"""
app.py — Web Ujian Online SMAN 1 Batangan
Flask Backend
"""
import os, uuid, random, string, io
from datetime import datetime, timezone
from flask import Flask, request, jsonify, send_file, redirect
from flask_cors import CORS
from dotenv import load_dotenv

load_dotenv()

from db   import query, get_db, release_db
from auth import (require_auth, require_guru, require_admin,
                  exchange_code_for_token, get_google_userinfo,
                  upsert_user, create_token, decode_token)

app = Flask(__name__)
CORS(app, origins="*", supports_credentials=True)

FRONTEND_URL = os.environ.get("FRONTEND_URL", "http://localhost:3000")
APP_URL      = os.environ.get("APP_URL",      "http://localhost:5000")
CLIENT_ID    = os.environ.get("GOOGLE_CLIENT_ID", "")

def now(): return datetime.now(timezone.utc)
def gen_token(n=32):
    return ''.join(random.choices(string.ascii_letters + string.digits, k=n))

# ═══════════════════════════════════════════════════════════
#  HEALTH
# ═══════════════════════════════════════════════════════════
@app.route("/")
def health():
    DEV_MODE = os.environ.get("DEV_MODE","true").lower() == "true"
    return jsonify({"status": "ok", "service": "Ujian Online SMAN 1 Batangan", "dev_mode": DEV_MODE})

# ═══════════════════════════════════════════════════════════
#  AUTH — Google OAuth
# ═══════════════════════════════════════════════════════════
@app.route("/api/auth/google/url")
def google_auth_url():
    """Generate URL login Google — dipanggil frontend saat klik 'Masuk'."""
    import urllib.parse
    role = request.args.get("role", "siswa")
    redirect_uri = f"{APP_URL}/api/auth/google/callback"
    # Simpan role di state parameter agar bisa dibaca saat callback
    state = urllib.parse.quote(role)
    params = "&".join([
        f"client_id={CLIENT_ID}",
        "response_type=code",
        f"redirect_uri={urllib.parse.quote(redirect_uri, safe='')}",
        "scope=openid%20email%20profile",
        "access_type=offline",
        "prompt=select_account",
        f"state={state}",
    ])
    url = f"https://accounts.google.com/o/oauth2/v2/auth?{params}"
    return jsonify({"url": url})

@app.route("/api/auth/google/callback")
def google_callback():
    """Google redirect ke sini setelah user login."""
    code  = request.args.get("code")
    error = request.args.get("error")

    if error or not code:
        return redirect(f"{FRONTEND_URL}/?error=login_gagal")

    try:
        import urllib.parse
        redirect_uri     = f"{APP_URL}/api/auth/google/callback"
        requested_role   = urllib.parse.unquote(request.args.get("state","siswa"))
        token_data       = exchange_code_for_token(code, redirect_uri)
        google_info      = get_google_userinfo(token_data["access_token"])
        result           = upsert_user(google_info, requested_role)
        token            = result["token"]
        role             = result["user"]["role"]

        # Redirect ke halaman sesuai role (struktur baru)
        # Selalu redirect ke index.html — biarkan frontend fetch user lalu redirect
        return redirect(f"{FRONTEND_URL}/index.html?token={token}")
    except ValueError as e:
        return redirect(f"{FRONTEND_URL}/?error={str(e)}")
    except Exception as e:
        return redirect(f"{FRONTEND_URL}/?error=server_error")

@app.route("/api/auth/me")
@require_auth
def me():
    """Ambil data user yang sedang login."""
    user = query("SELECT * FROM users WHERE id=%s", (request.user_id,), fetch="one")
    if not user:
        return jsonify({"error": "User tidak ditemukan"}), 404
    u = dict(user)
    u.pop("google_id", None)  # jangan expose google_id
    return jsonify(u)

# ═══════════════════════════════════════════════════════════
#  ADMIN — Kelola User
# ═══════════════════════════════════════════════════════════
@app.route("/api/admin/users", methods=["GET"])
@require_admin
def list_users():
    role = request.args.get("role")
    if role:
        users = query("SELECT * FROM users WHERE role=%s ORDER BY name", (role,))
    else:
        users = query("SELECT * FROM users ORDER BY role, name")
    return jsonify([dict(u) for u in users])

@app.route("/api/admin/users/<user_id>", methods=["PATCH"])
@require_admin
def update_user(user_id):
    data = request.json or {}
    if "role" in data:
        query("UPDATE users SET role=%s WHERE id=%s",
              (data["role"], user_id), fetch="none")
    if "class_id" in data:
        query("UPDATE users SET class_id=%s WHERE id=%s",
              (data["class_id"], user_id), fetch="none")
    if "is_active" in data:
        query("UPDATE users SET is_active=%s WHERE id=%s",
              (data["is_active"], user_id), fetch="none")
    user = query("SELECT * FROM users WHERE id=%s", (user_id,), fetch="one")
    return jsonify(dict(user))

# ═══════════════════════════════════════════════════════════
#  DEV: Login Dummy (hanya aktif saat DEV_MODE=true)
# ═══════════════════════════════════════════════════════════
@app.route("/api/dev/login-as", methods=["POST"])
def dev_login_as():
    """Login sebagai user tertentu tanpa Google OAuth — hanya untuk testing."""
    DEV_MODE = os.environ.get("DEV_MODE","true").lower() == "true"
    if not DEV_MODE:
        return jsonify({"error": "Hanya tersedia di mode development"}), 403
    data = request.json or {}
    email = data.get("email","").strip()
    if not email:
        return jsonify({"error": "Email wajib diisi"}), 400
    user = query("SELECT * FROM users WHERE email=%s", (email,), fetch="one")
    if not user:
        return jsonify({"error": f"User {email} tidak ditemukan"}), 404
    from auth import create_token
    token = create_token(str(user["id"]), user["role"])
    return jsonify({"token": token, "user": dict(user)})

# ═══════════════════════════════════════════════════════════
#  KELAS & MATA PELAJARAN
# ═══════════════════════════════════════════════════════════
@app.route("/api/classes", methods=["GET"])
@require_auth
def get_classes():
    rows = query("SELECT * FROM classes ORDER BY grade, LENGTH(id), id")
    return jsonify([dict(r) for r in rows])

@app.route("/api/subjects", methods=["GET"])
@require_guru
def get_subjects():
    """Ambil mapel milik guru ini saja."""
    rows = query("""
        SELECT * FROM subjects
        WHERE teacher_id=%s OR teacher_id IS NULL
        ORDER BY name
    """, (request.user_id,))
    return jsonify([dict(r) for r in rows])

@app.route("/api/admin/subjects", methods=["GET"])
@require_admin
def admin_get_all_subjects():
    """Admin lihat semua mapel dari semua guru."""
    rows = query("""
        SELECT s.*, u.name as teacher_name
        FROM subjects s
        LEFT JOIN users u ON u.id = s.teacher_id
        ORDER BY u.name, s.name
    """)
    return jsonify([dict(r) for r in rows])

@app.route("/api/subjects", methods=["POST"])
@require_guru
def create_subject():
    data = request.json or {}
    name = data.get("name","").strip()
    if not name:
        return jsonify({"error": "Nama mapel wajib diisi"}), 400
    existing = query("SELECT id FROM subjects WHERE LOWER(name)=LOWER(%s)", (name,), fetch="one")
    if existing:
        return jsonify({"error": "Mapel sudah ada"}), 409
    sub = query("""
        INSERT INTO subjects (id, name, teacher_id) VALUES (%s, %s, %s) RETURNING *
    """, (str(uuid.uuid4()), name, request.user_id), fetch="one")
    return jsonify(dict(sub)), 201

@app.route("/api/subjects/<subject_id>", methods=["PATCH"])
@require_guru
def update_subject(subject_id):
    data = request.json or {}
    name = data.get("name","").strip()
    if not name:
        return jsonify({"error": "Nama mapel wajib diisi"}), 400
    query("UPDATE subjects SET name=%s WHERE id=%s", (name, subject_id), fetch="none")
    sub = query("SELECT * FROM subjects WHERE id=%s", (subject_id,), fetch="one")
    return jsonify(dict(sub))

@app.route("/api/subjects/<subject_id>", methods=["DELETE"])
@require_guru
def delete_subject(subject_id):
    sub = query("SELECT * FROM subjects WHERE id=%s", (subject_id,), fetch="one")
    if not sub:
        return jsonify({"error": "Mapel tidak ditemukan"}), 404
    if str(sub.get("teacher_id","")) != request.user_id and request.user_role != "admin":
        return jsonify({"error": "Bukan mapel kamu"}), 403
    used = query("SELECT id FROM exams WHERE subject_id=%s LIMIT 1", (subject_id,), fetch="one")
    if used:
        return jsonify({"error": "Mapel masih dipakai di ujian"}), 400
    query("DELETE FROM subjects WHERE id=%s", (subject_id,), fetch="none")
    return jsonify({"ok": True})





# ═══════════════════════════════════════════════════════════
#  UJIAN — CRUD
# ═══════════════════════════════════════════════════════════
@app.route("/api/exams", methods=["GET"])
@require_guru
def list_exams():
    if request.user_role == "admin":
        rows = query("""
            SELECT e.*, u.name as teacher_name, s.name as subject_name
            FROM exams e
            LEFT JOIN users    u ON e.teacher_id  = u.id
            LEFT JOIN subjects s ON e.subject_id  = s.id
            ORDER BY e.created_at DESC
        """)
    else:
        rows = query("""
            SELECT e.*, u.name as teacher_name, s.name as subject_name
            FROM exams e
            LEFT JOIN users    u ON e.teacher_id = u.id
            LEFT JOIN subjects s ON e.subject_id = s.id
            WHERE e.teacher_id = %s
            ORDER BY e.created_at DESC
        """, (request.user_id,))
    result = []
    for r in rows:
        d = dict(r)
        # Hitung jumlah soal
        q = query("SELECT COUNT(*) as c FROM questions WHERE exam_id=%s",
                  (d["id"],), fetch="one")
        d["question_count"] = q["c"] if q else 0
        # Kelas peserta
        cls = query("""
            SELECT c.id, c.name FROM exam_classes ec
            JOIN classes c ON ec.class_id=c.id WHERE ec.exam_id=%s
        """, (d["id"],))
        d["classes"] = [dict(c) for c in cls]
        result.append(d)
    return jsonify(result)

@app.route("/api/exams", methods=["POST"])
@require_guru
def create_exam():
    data = request.json or {}
    required = ["title", "start_at", "duration_minutes"]
    for f in required:
        if not data.get(f):
            return jsonify({"error": f"Field '{f}' wajib diisi"}), 400

    eid = str(uuid.uuid4())
    query("""
        INSERT INTO exams
          (id, title, subject_id, teacher_id, start_at, duration_minutes,
           randomize_questions, randomize_options, show_result_after,
           show_key_after, instructions, status)
        VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,'draft')
    """, (
        eid,
        data["title"],
        data.get("subject_id"),
        request.user_id,
        data["start_at"],
        data["duration_minutes"],
        data.get("randomize_questions", True),
        data.get("randomize_options",   True),
        data.get("show_result_after",   True),
        data.get("show_key_after",      False),
        data.get("instructions", ""),
    ), fetch="none")

    # Daftarkan kelas peserta
    for class_id in data.get("class_ids", []):
        query("INSERT INTO exam_classes (exam_id,class_id) VALUES (%s,%s) ON CONFLICT DO NOTHING",
              (eid, class_id), fetch="none")

    exam = query("SELECT * FROM exams WHERE id=%s", (eid,), fetch="one")
    return jsonify(dict(exam)), 201

@app.route("/api/exams/<exam_id>", methods=["GET"])
@require_guru
def get_exam(exam_id):
    exam = query("""
        SELECT e.*, u.name as teacher_name, s.name as subject_name
        FROM exams e
        LEFT JOIN users    u ON e.teacher_id = u.id
        LEFT JOIN subjects s ON e.subject_id = s.id
        WHERE e.id=%s
    """, (exam_id,), fetch="one")
    if not exam:
        return jsonify({"error": "Ujian tidak ditemukan"}), 404
    d = dict(exam)

    # Ambil soal
    questions = query(
        "SELECT * FROM questions WHERE exam_id=%s ORDER BY order_num",
        (exam_id,)
    )

    # Ambil semua options sekaligus dengan JOIN
    opts_map = {}
    if questions:
        from psycopg2.extras import execute_values
        opts = query("""
            SELECT o.* FROM options o
            JOIN questions q ON o.question_id = q.id
            WHERE q.exam_id = %s
            ORDER BY o.label
        """, (exam_id,))
        for o in opts:
            qid = str(o["question_id"])
            if qid not in opts_map:
                opts_map[qid] = []
            opts_map[qid].append(dict(o))

    d["questions"] = []
    for q in questions:
        qd = dict(q)
        qd["options"] = opts_map.get(str(q["id"]), [])
        d["questions"].append(qd)

    # Kelas peserta
    cls = query("""
        SELECT c.* FROM exam_classes ec
        JOIN classes c ON ec.class_id=c.id WHERE ec.exam_id=%s
    """, (exam_id,))
    d["classes"] = [dict(c) for c in cls]
    return jsonify(d)

@app.route("/api/exams/<exam_id>", methods=["PATCH"])
@require_guru
def update_exam(exam_id):
    data = request.json or {}
    fields = ["title","start_at","duration_minutes","randomize_questions",
              "randomize_options","show_result_after","show_key_after",
              "instructions","status","subject_id"]
    for f in fields:
        if f in data:
            query(f"UPDATE exams SET {f}=%s WHERE id=%s",
                  (data[f], exam_id), fetch="none")
    if "class_ids" in data:
        query("DELETE FROM exam_classes WHERE exam_id=%s", (exam_id,), fetch="none")
        for cid in data["class_ids"]:
            query("INSERT INTO exam_classes VALUES (%s,%s) ON CONFLICT DO NOTHING",
                  (exam_id, cid), fetch="none")
    return jsonify({"ok": True})

@app.route("/api/exams/<exam_id>", methods=["DELETE"])
@require_guru
def delete_exam(exam_id):
    query("DELETE FROM exams WHERE id=%s", (exam_id,), fetch="none")
    return jsonify({"ok": True})

# ═══════════════════════════════════════════════════════════
#  SOAL
# ═══════════════════════════════════════════════════════════
@app.route("/api/exams/<exam_id>/questions", methods=["POST"])
@require_guru
def add_question(exam_id):
    data = request.json or {}
    content = data.get("content","").strip()
    if not content:
        return jsonify({"error": "Isi soal wajib diisi"}), 400
    options = data.get("options", [])
    if len(options) < 2:
        return jsonify({"error": "Minimal 2 pilihan jawaban"}), 400
    if not any(o.get("is_correct") for o in options):
        return jsonify({"error": "Harus ada 1 kunci jawaban"}), 400

    # Urutan soal berikutnya
    last = query("SELECT MAX(order_num) as m FROM questions WHERE exam_id=%s",
                 (exam_id,), fetch="one")
    order_num = (last["m"] or 0) + 1

    qid = str(uuid.uuid4())
    query("""
        INSERT INTO questions (id, exam_id, content, image_url, order_num, score)
        VALUES (%s, %s, %s, %s, %s, %s)
    """, (qid, exam_id, content, data.get("image_url"), order_num,
          data.get("score", 1)), fetch="none")

    for opt in options:
        query("""
            INSERT INTO options (id, question_id, label, content, is_correct)
            VALUES (%s, %s, %s, %s, %s)
        """, (str(uuid.uuid4()), qid, opt["label"], opt["content"],
              opt.get("is_correct", False)), fetch="none")

    q = query("SELECT * FROM questions WHERE id=%s", (qid,), fetch="one")
    opts = query("SELECT * FROM options WHERE question_id=%s ORDER BY label", (qid,))
    result = dict(q)
    result["options"] = [dict(o) for o in opts]
    return jsonify(result), 201

@app.route("/api/questions/<question_id>", methods=["DELETE"])
@require_guru
def delete_question(question_id):
    query("DELETE FROM questions WHERE id=%s", (question_id,), fetch="none")
    return jsonify({"ok": True})

@app.route("/api/questions/<question_id>", methods=["PATCH"])
@require_guru
def update_question(question_id):
    data = request.json or {}
    if "content" in data:
        query("UPDATE questions SET content=%s WHERE id=%s",
              (data["content"], question_id), fetch="none")
    if "options" in data:
        query("DELETE FROM options WHERE question_id=%s", (question_id,), fetch="none")
        for opt in data["options"]:
            query("""
                INSERT INTO options (id, question_id, label, content, is_correct)
                VALUES (%s, %s, %s, %s, %s)
            """, (str(uuid.uuid4()), question_id, opt["label"],
                  opt["content"], opt.get("is_correct", False)), fetch="none")
    return jsonify({"ok": True})

# ═══════════════════════════════════════════════════════════
#  IMPORT SOAL dari CSV / Excel / Word
# ═══════════════════════════════════════════════════════════
@app.route("/api/exams/<exam_id>/import", methods=["POST"])
@require_guru
def import_questions(exam_id):
    if "file" not in request.files:
        return jsonify({"error": "File tidak ditemukan"}), 400
    f    = request.files["file"]
    ext  = f.filename.rsplit(".",1)[-1].lower()
    rows = []

    try:
        if ext == "csv":
            import csv, io as _io
            raw = f.read()
            # Handle BOM (utf-8-sig) dan Windows line ending
            content = raw.decode("utf-8-sig", errors="ignore").replace("\r\n","\n").replace("\r","\n")
            reader  = csv.DictReader(_io.StringIO(content))
            for row in reader:
                rows.append(row)

        elif ext in ("xlsx", "xls"):
            from openpyxl import load_workbook
            wb   = load_workbook(f, read_only=True)
            ws   = wb.active
            hdrs = [str(c.value).strip().lower() if c.value else "" for c in next(ws.iter_rows())]
            for row in ws.iter_rows(min_row=2, values_only=True):
                rows.append({hdrs[i]: (str(v).strip() if v is not None else "")
                             for i, v in enumerate(row)})

        elif ext == "docx":
            # Simpan ke tempfile lalu parse
            import tempfile, os
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                f.save(tmp.name)
                tmppath = tmp.name
            try:
                rows = _parse_docx_new(tmppath)
            finally:
                os.unlink(tmppath)
            return _save_imported_questions(exam_id, rows)

        elif ext == "doc":
            # Word lama (.doc) — pakai antiword
            import tempfile, os
            with tempfile.NamedTemporaryFile(delete=False, suffix=".doc") as tmp:
                f.save(tmp.name)
                tmppath = tmp.name
            try:
                rows = _parse_doc_old(tmppath)
            finally:
                os.unlink(tmppath)
            return _save_imported_questions(exam_id, rows)

        else:
            return jsonify({"error": "Format tidak didukung. Gunakan CSV, XLSX, DOCX, atau DOC"}), 400

    except Exception as e:
        return jsonify({"error": f"Gagal membaca file: {str(e)}"}), 400

    return _save_imported_questions(exam_id, rows)

def _parse_doc_old(filepath):
    """Baca .doc lama (Word 97-2003) menggunakan antiword."""
    import subprocess
    result = subprocess.run(
        ['antiword', filepath],
        capture_output=True, text=True, timeout=30
    )
    if result.returncode != 0:
        raise Exception("Gagal membaca file .doc")
    return _parse_word_text(result.stdout)

def _parse_docx_new(filepath):
    """
    Parser universal .docx — handle semua format soal Indonesia:
    - Format 1: "1. Soal.... a. Pilihan A" lalu b/c/d/e di paragraf berikutnya
    - Format 2: "1. Soal" di paragraf sendiri, pilihan di paragraf berikutnya
    - Format 3: Soal + semua pilihan dalam 1 paragraf panjang
    - Deteksi kunci jawaban otomatis dari teks TEBAL (bold)
    """
    from docx import Document
    import re

    def _valid(q):
        return q.get('soal','').strip() and sum(1 for k in 'abcde' if q.get(k,'').strip()) >= 2

    doc = Document(filepath)
    questions = []
    current = None
    SKIP = ['SOAL PILIHAN','PILIHLAH','NAMA','KELAS','ESSAY','URAIAN',
            'JAWABLAH','BERILAH','PETUNJUK','LEMBAR','SATUAN','KURIKULUM',
            'MATA PELAJARAN','SEMESTER','SEKOLAH','PENILAIAN','TAHUN']

    for p in doc.paragraphs:
        text = re.sub(r'[ ​ ]+', ' ', p.text).strip()
        if not text: continue
        if any(s in text.upper()[:60] for s in SKIP): continue

        is_bold_para = any(r.bold and r.text.strip() for r in p.runs)

        # FORMAT 1: "1. Soal.... a. Pilihan A" inline
        m1 = re.match(r'^(\d+)\s*[.)]\s*(.+?)\s+a\.\s+(.+)$', text)
        if m1 and 1 <= int(m1.group(1)) <= 200:
            if current and _valid(current):
                questions.append(current)
            current = {
                'soal': m1.group(2).strip(),
                'a': m1.group(3).strip(),
                'b':'','c':'','d':'','e':'', 'jawaban':''
            }
            if is_bold_para and not current['jawaban']:
                current['jawaban'] = 'A'
            continue

        # FORMAT 2: "1. Soal" tanpa pilihan
        m2 = re.match(r'^(\d+)\s*[.)]\s*(.+)$', text)
        if m2 and 1 <= int(m2.group(1)) <= 200 and not re.match(r'^[a-eA-E][.)]\s', text):
            if current and _valid(current):
                questions.append(current)
            current = {'soal': m2.group(2).strip(), 'a':'','b':'','c':'','d':'','e':'', 'jawaban':''}
            continue

        if not current: continue

        # Pilihan a-e
        m3 = re.match(r'^([a-eA-E])\s*[.)]\s*(.+)$', text)
        if m3:
            label = m3.group(1).lower()
            current[label] = m3.group(2).strip()
            if is_bold_para and not current['jawaban']:
                current['jawaban'] = label.upper()
            continue

        # Lanjutan soal
        if not any(current.get(k,'').strip() for k in 'abcde'):
            current['soal'] += ' ' + text

    if current and _valid(current):
        questions.append(current)
    return questions


def _parse_word_text(text):
    """
    Parser fleksibel soal ujian Indonesia.
    Support: nomor 1./1), pilihan a./A./a)/A),
             format 1 kolom & 2 kolom, wrap teks soal.
    """
    import re
    questions = []
    current = None
    lines = text.split('\n')
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        # Nomor soal: '   1. teks' atau '  10. teks'
        m = re.match(r'^\s{1,}(\d+)[.)\s]\s*(.+)', line)
        if m and 1 <= int(m.group(1)) <= 200:
            if current and sum(1 for k in 'abcde' if current.get(k)):
                questions.append(current)
            current = {'no':int(m.group(1)),'soal':m.group(2).strip(),
                       'a':'','b':'','c':'','d':'','e':'','jawaban':''}
            continue
        if not current:
            continue
        # Pilihan 2 kolom
        m2 = re.match(r'^\s+([a-eA-E])[.)\s]\s{1,8}(.+?)\s{4,}([a-eA-E])[.)\s]\s{1,8}(.+)', line)
        if m2:
            current[m2.group(1).lower()] = m2.group(2).strip()
            current[m2.group(3).lower()] = m2.group(4).strip()
            continue
        # Pilihan 1 kolom
        m3 = re.match(r'^\s+([a-eA-E])[.)\s]\s+(.+)', line)
        if m3:
            current[m3.group(1).lower()] = m3.group(2).strip()
            continue
        # Kunci jawaban
        m4 = re.match(r'(?:jawaban|kunci|answer)[:\s]+([a-eA-E])', stripped, re.IGNORECASE)
        if m4:
            current['jawaban'] = m4.group(1).upper()
            continue
        # Lanjutan soal
        if not any(current.get(k) for k in 'abcde'):
            current['soal'] += ' ' + stripped
    if current and sum(1 for k in 'abcde' if current.get(k)):
        questions.append(current)
    return questions


def _save_imported_questions(exam_id, rows):
    """Simpan soal ke database. rows = list of dict dari parser."""
    saved = 0
    errors = []
    last = query("SELECT MAX(order_num) as m FROM questions WHERE exam_id=%s",
                 (exam_id,), fetch="one")
    order_num = int(last["m"] or 0) + 1

    for i, row in enumerate(rows, 1):
        try:
            soal = (row.get('soal') or row.get('pertanyaan') or
                    row.get('question') or row.get('content') or '').strip()
            if not soal or len(soal) < 3:
                continue

            jawaban = (row.get('jawaban') or row.get('kunci') or
                       row.get('answer') or '').strip().upper()

            # Kumpulkan pilihan (support huruf besar & kecil)
            pilihan = {}
            for label in ['A','B','C','D','E']:
                val = (row.get(label) or row.get(label.lower()) or '').strip()
                if val:
                    pilihan[label] = val

            if len(pilihan) < 2:
                errors.append(f"Soal {i}: kurang dari 2 pilihan")
                continue

            if not jawaban or jawaban not in pilihan:
                jawaban = ''  # kosong = guru set manual

            qid = str(uuid.uuid4())
            query("INSERT INTO questions (id,exam_id,content,order_num) VALUES (%s,%s,%s,%s)",
                  (qid, exam_id, soal, order_num), fetch="none")

            for label, val in pilihan.items():
                query("INSERT INTO options (id,question_id,label,content,is_correct) VALUES (%s,%s,%s,%s,%s)",
                      (str(uuid.uuid4()), qid, label, val, label==jawaban), fetch="none")

            order_num += 1
            saved += 1
        except Exception as e:
            errors.append(f"Soal {i}: {str(e)}")

    return jsonify({"saved": saved, "total": len(rows), "errors": errors})


@app.route("/api/student/exams", methods=["GET"])
@require_auth
def student_exams():
    """Ujian yang tersedia untuk siswa ini."""
    # Cek kelas siswa
    user = query("SELECT * FROM users WHERE id=%s", (request.user_id,), fetch="one")
    if not user or not user["class_id"]:
        return jsonify({"error": "Kelas siswa belum diatur"}), 400

    rows = query("""
        SELECT e.*, s.name as subject_name, u.name as teacher_name
        FROM exams e
        JOIN exam_classes ec ON ec.exam_id = e.id
        LEFT JOIN subjects s ON e.subject_id = s.id
        LEFT JOIN users    u ON e.teacher_id = u.id
        WHERE ec.class_id = %s
          AND e.status IN ('published','ongoing','finished')
        ORDER BY e.start_at DESC
    """, (user["class_id"],))

    result = []
    for r in rows:
        d = dict(r)
        # Cek apakah siswa sudah punya sesi
        sess = query("""
            SELECT id, submitted_at FROM exam_sessions
            WHERE exam_id=%s AND student_id=%s
        """, (d["id"], request.user_id), fetch="one")
        d["session_id"]   = str(sess["id"]) if sess else None
        d["submitted_at"] = sess["submitted_at"] if sess else None
        d["has_submitted"]= sess is not None and sess["submitted_at"] is not None
        result.append(d)

    return jsonify(result)

@app.route("/api/student/exams/<exam_id>/start", methods=["POST"])
@require_auth
def start_exam(exam_id):
    """Siswa mulai ujian — buat sesi baru."""
    data       = request.json or {}
    device_key = data.get("device_key","").strip()
    if not device_key:
        return jsonify({"error": "Device key wajib"}), 400

    # Cek ujian ada dan sudah published/ongoing
    exam = query("SELECT * FROM exams WHERE id=%s", (exam_id,), fetch="one")
    if not exam:
        return jsonify({"error": "Ujian tidak ditemukan"}), 404
    if exam["status"] not in ("published","ongoing"):
        return jsonify({"error": "Ujian belum dibuka"}), 400

    # Cek siswa terdaftar di kelas ujian ini
    user = query("SELECT * FROM users WHERE id=%s", (request.user_id,), fetch="one")
    eligible = query("""
        SELECT 1 FROM exam_classes WHERE exam_id=%s AND class_id=%s
    """, (exam_id, user["class_id"]), fetch="one")
    if not eligible:
        return jsonify({"error": "Kamu tidak terdaftar untuk ujian ini"}), 403

    # Cek sesi lama
    existing = query("""
        SELECT * FROM exam_sessions WHERE exam_id=%s AND student_id=%s
    """, (exam_id, request.user_id), fetch="one")
    if existing:
        if existing["submitted_at"]:
            return jsonify({"error": "Kamu sudah mengumpulkan ujian ini"}), 409
        # Lanjut sesi lama
        return _build_exam_session(existing, exam)

    # Cek device sudah dipakai
    dev_used = query("""
        SELECT 1 FROM exam_sessions WHERE exam_id=%s AND device_key=%s
    """, (exam_id, device_key), fetch="one")
    if dev_used:
        return jsonify({"error": "Perangkat ini sudah digunakan untuk ujian ini"}), 409

    # Buat sesi baru
    sess_id = str(uuid.uuid4())
    token   = gen_token(48)
    query("""
        INSERT INTO exam_sessions
          (id, exam_id, student_id, token, device_key, ip_address)
        VALUES (%s, %s, %s, %s, %s, %s)
    """, (sess_id, exam_id, request.user_id, token,
          device_key, request.remote_addr), fetch="none")

    # Update status ujian jadi ongoing
    query("UPDATE exams SET status='ongoing' WHERE id=%s AND status='published'",
          (exam_id,), fetch="none")

    sess = query("SELECT * FROM exam_sessions WHERE id=%s", (sess_id,), fetch="one")
    return _build_exam_session(sess, exam)

def _build_exam_session(sess, exam):
    """Bangun response sesi ujian untuk siswa."""
    # Ambil soal (dengan/tanpa acak)
    questions = query("""
        SELECT * FROM questions WHERE exam_id=%s ORDER BY order_num
    """, (exam["id"],))

    q_list = list(questions)
    if exam["randomize_questions"]:
        random.shuffle(q_list)

    result_questions = []
    for q in q_list:
        opts = query("SELECT * FROM options WHERE question_id=%s ORDER BY label",
                     (q["id"],))
        opt_list = list(opts)
        if exam["randomize_options"]:
            random.shuffle(opt_list)
        # JANGAN kirim is_correct ke siswa!
        result_questions.append({
            "id":      str(q["id"]),
            "content": q["content"],
            "image_url": q["image_url"],
            "options": [{"id": str(o["id"]), "label": o["label"],
                         "content": o["content"]} for o in opt_list]
        })

    # Jawaban yang sudah diisi
    answers = query("""
        SELECT question_id, option_id FROM answers WHERE session_id=%s
    """, (sess["id"],))
    saved_answers = {str(a["question_id"]): str(a["option_id"])
                     for a in answers if a["option_id"]}

    # Hitung sisa waktu
    from datetime import timedelta
    start   = sess["started_at"]
    end_at  = start + timedelta(minutes=int(exam["duration_minutes"]))
    remaining = max(0, (end_at - datetime.now(timezone.utc)).total_seconds())

    return jsonify({
        "session_id":   str(sess["id"]),
        "token":        sess["token"],
        "exam_title":   exam["title"],
        "duration":     exam["duration_minutes"],
        "remaining_seconds": int(remaining),
        "end_at":       end_at.isoformat(),
        "questions":    result_questions,
        "saved_answers": saved_answers,
    })

@app.route("/api/student/sessions/<session_id>/answer", methods=["POST"])
@require_auth
def save_answer(session_id):
    """Auto-save jawaban siswa (dipanggil tiap 30 detik atau saat ganti soal)."""
    data        = request.json or {}
    question_id = data.get("question_id")
    option_id   = data.get("option_id")  # boleh null (belum dijawab)

    sess = query("SELECT * FROM exam_sessions WHERE id=%s AND student_id=%s",
                 (session_id, request.user_id), fetch="one")
    if not sess:
        return jsonify({"error": "Sesi tidak ditemukan"}), 404
    if sess["submitted_at"]:
        return jsonify({"error": "Ujian sudah dikumpulkan"}), 400

    # Upsert jawaban
    existing = query("SELECT id FROM answers WHERE session_id=%s AND question_id=%s",
                     (session_id, question_id), fetch="one")
    if existing:
        query("UPDATE answers SET option_id=%s, answered_at=NOW() WHERE id=%s",
              (option_id, existing["id"]), fetch="none")
    else:
        query("""
            INSERT INTO answers (id, session_id, question_id, option_id)
            VALUES (%s, %s, %s, %s)
        """, (str(uuid.uuid4()), session_id, question_id, option_id), fetch="none")

    # Update aktivitas terakhir
    query("UPDATE exam_sessions SET last_activity=NOW() WHERE id=%s",
          (session_id,), fetch="none")

    return jsonify({"ok": True})

@app.route("/api/student/sessions/<session_id>/violation", methods=["POST"])
@require_auth
def record_violation(session_id):
    """Catat pelanggaran (pindah tab)."""
    query("""
        UPDATE exam_sessions
        SET tab_violations = tab_violations + 1
        WHERE id=%s AND student_id=%s
    """, (session_id, request.user_id), fetch="none")
    sess = query("SELECT tab_violations FROM exam_sessions WHERE id=%s",
                 (session_id,), fetch="one")
    return jsonify({"tab_violations": sess["tab_violations"] if sess else 0})

@app.route("/api/student/sessions/<session_id>/submit", methods=["POST"])
@require_auth
def submit_exam(session_id):
    """Siswa kumpulkan ujian."""
    data        = request.json or {}
    auto_submit = data.get("auto_submit", False)

    sess = query("SELECT * FROM exam_sessions WHERE id=%s AND student_id=%s",
                 (session_id, request.user_id), fetch="one")
    if not sess:
        return jsonify({"error": "Sesi tidak ditemukan"}), 404
    if sess["submitted_at"]:
        return jsonify({"error": "Sudah dikumpulkan"}), 400

    query("""
        UPDATE exam_sessions
        SET submitted_at=NOW(), auto_submitted=%s
        WHERE id=%s
    """, (auto_submit, session_id), fetch="none")

    # Hitung nilai langsung di Python (tanpa stored procedure)
    exam_data = query("SELECT * FROM exams WHERE id=%s", (sess["exam_id"],), fetch="one")

    # Ambil semua soal ujian
    total_q = query("SELECT COUNT(*) as n FROM questions WHERE exam_id=%s",
                    (sess["exam_id"],), fetch="one")["n"]

    # Hitung jawaban benar
    correct = query("""
        SELECT COUNT(*) as n FROM answers a
        JOIN options o ON o.id = a.option_id
        WHERE a.session_id=%s AND o.is_correct=true
    """, (session_id,), fetch="one")["n"]

    wrong = query("""
        SELECT COUNT(*) as n FROM answers a
        JOIN options o ON o.id = a.option_id
        WHERE a.session_id=%s AND o.is_correct=false
    """, (session_id,), fetch="one")["n"]

    empty = total_q - correct - wrong

    # Hitung skor dengan formula dari ujian
    # score_formula: 'standard' (benar/total*100), atau custom
    score_per_correct = float(exam_data.get("score_per_correct") or (100.0/total_q if total_q else 0))
    score = round(correct * score_per_correct, 2)

    # Simpan hasil
    existing = query("SELECT id FROM results WHERE session_id=%s", (session_id,), fetch="one")
    if existing:
        query("""UPDATE results SET score=%s, correct_count=%s, wrong_count=%s, empty_count=%s
                 WHERE session_id=%s""",
              (score, correct, wrong, empty, session_id), fetch="none")
    else:
        query("""INSERT INTO results (id, session_id, score, correct_count, wrong_count, empty_count)
                 VALUES (%s, %s, %s, %s, %s, %s)""",
              (str(uuid.uuid4()), session_id, score, correct, wrong, empty), fetch="none")

    result = query("SELECT * FROM results WHERE session_id=%s", (session_id,), fetch="one")
    exam   = query("SELECT show_result_after FROM exams WHERE id=%s",
                   (sess["exam_id"],), fetch="one")

    resp = {"ok": True, "submitted_at": datetime.now(timezone.utc).isoformat()}
    if exam and exam["show_result_after"] and result:
        resp["score"]          = float(result["score"])
        resp["correct_count"]  = result["correct_count"]
        resp["total_questions"]= result["total_questions"]

    return jsonify(resp)

# ═══════════════════════════════════════════════════════════
#  PENGAWAS — Live Monitor (guru)
# ═══════════════════════════════════════════════════════════
@app.route("/api/exams/<exam_id>/monitor", methods=["GET"])
@require_guru
def monitor_exam(exam_id):
    """Data live untuk halaman pengawas."""
    exam = query("SELECT * FROM exams WHERE id=%s", (exam_id,), fetch="one")
    if not exam:
        return jsonify({"error": "Ujian tidak ditemukan"}), 404

    total_q = query("SELECT COUNT(*) as c FROM questions WHERE exam_id=%s",
                    (exam_id,), fetch="one")["c"]

    # Semua siswa yang terdaftar
    students = query("""
        SELECT u.id, u.name, u.avatar_url, c.name as class_name
        FROM users u
        JOIN classes c ON u.class_id = c.id
        JOIN exam_classes ec ON ec.class_id = c.id
        WHERE ec.exam_id=%s AND u.role='siswa'
        ORDER BY c.grade, LENGTH(c.id), c.id, u.name
    """, (exam_id,))

    # Sesi yang ada
    sessions = query("""
        SELECT es.*, r.score
        FROM exam_sessions es
        LEFT JOIN results r ON r.session_id = es.id
        WHERE es.exam_id=%s
    """, (exam_id,))
    sess_map = {str(s["student_id"]): dict(s) for s in sessions}

    result = []
    for stu in students:
        sid  = str(stu["id"])
        sess = sess_map.get(sid)

        answered = 0
        if sess:
            ans = query("""
                SELECT COUNT(*) as c FROM answers
                WHERE session_id=%s AND option_id IS NOT NULL
            """, (sess["id"],), fetch="one")
            answered = ans["c"] if ans else 0

        result.append({
            "student_id":     sid,
            "session_id":     sess["id"] if sess else None,
            "name":           stu["name"],
            "avatar_url":     stu["avatar_url"],
            "class_name":     stu["class_name"],
            "status":         ("submitted" if sess and sess["submitted_at"]
                               else "online" if sess else "offline"),
            "answered":       answered,
            "total_q":        total_q,
            "tab_violations": sess["tab_violations"] if sess else 0,
            "submitted_at":   sess["submitted_at"].isoformat() if sess and sess["submitted_at"] else None,
            "score":          float(sess["score"]) if sess and sess.get("score") else None,
        })

    submitted = sum(1 for r in result if r["status"] == "submitted")
    online    = sum(1 for r in result if r["status"] == "online")
    offline   = sum(1 for r in result if r["status"] == "offline")

    return jsonify({
        "exam":      dict(exam),
        "summary":   {"submitted": submitted, "online": online,
                      "offline": offline, "total": len(result)},
        "students":  result,
    })

@app.route("/api/exams/<exam_id>/finish", methods=["POST"])
@require_guru
def finish_exam(exam_id):
    """Guru tutup ujian secara paksa."""
    query("UPDATE exams SET status='finished' WHERE id=%s", (exam_id,), fetch="none")
    return jsonify({"ok": True})

@app.route("/api/exams/<exam_id>/publish-results", methods=["POST"])
@require_guru
def publish_results(exam_id):
    """Guru publish nilai agar siswa bisa lihat."""
    query("UPDATE results SET published=TRUE WHERE exam_id=%s", (exam_id,), fetch="none")
    return jsonify({"ok": True})

# ═══════════════════════════════════════════════════════════
#  NILAI & REKAP
# ═══════════════════════════════════════════════════════════
@app.route("/api/exams/<exam_id>/results", methods=["GET"])
@require_guru
def exam_results(exam_id):
    # Ambil semua siswa yang punya sesi di ujian ini + hasil nilainya
    rows = query("""
        SELECT
            u.id as student_id, u.name, u.nisn,
            c.name as class_name,
            es.submitted_at, es.tab_violations,
            r.score, r.correct_count, r.wrong_count, r.empty_count
        FROM exam_sessions es
        JOIN users u ON u.id = es.student_id
        LEFT JOIN classes c ON c.id = u.class_id
        LEFT JOIN results r ON r.session_id = es.id
        WHERE es.exam_id = %s
        ORDER BY c.grade, LENGTH(c.id), c.id, u.name
    """, (exam_id,))

    # Hitung summary
    all_rows = [dict(r) for r in rows]
    submitted = [r for r in all_rows if r.get('submitted_at')]
    scores = [float(r['score']) for r in submitted if r.get('score') is not None]

    # Distribusi nilai
    dist = {'a':0,'b':0,'c':0,'d':0}
    for s in scores:
        if s >= 90: dist['a'] += 1
        elif s >= 75: dist['b'] += 1
        elif s >= 60: dist['c'] += 1
        else: dist['d'] += 1

    # Analisis per soal
    questions = query("""
        SELECT q.id, q.content,
               COUNT(a.id) FILTER (WHERE o.is_correct=true) as correct,
               COUNT(a.id) as total
        FROM questions q
        LEFT JOIN answers a ON a.question_id = q.id
            AND a.session_id IN (SELECT id FROM exam_sessions WHERE exam_id=%s)
        LEFT JOIN options o ON o.id = a.option_id
        WHERE q.exam_id=%s
        GROUP BY q.id, q.content
        ORDER BY q.order_num
    """, (exam_id, exam_id))

    q_stats = []
    for q in questions:
        total = q['total'] or 1
        pct = round((q['correct'] or 0) / total * 100, 1)
        q_stats.append({'correct': q['correct'] or 0, 'total': total, 'pct': pct})

    summary = {
        'total_students': len(all_rows),
        'submitted': len(submitted),
        'avg_score': round(sum(scores)/len(scores), 1) if scores else None,
        'pass_rate': round(len([s for s in scores if s >= 75])/len(scores)*100, 1) if scores else 0,
    }

    return jsonify({
        'results': all_rows,
        'summary': summary,
        'score_distribution': dist,
        'question_stats': q_stats,
    })

@app.route("/api/exams/<exam_id>/student/<student_id>/detail", methods=["GET"])
@require_guru
def student_exam_detail(exam_id, student_id):
    """Detail jawaban siswa via exam_id + student_id (lebih mudah dari frontend)."""
    sess = query("""
        SELECT id FROM exam_sessions
        WHERE exam_id=%s AND student_id=%s
        LIMIT 1
    """, (exam_id, student_id), fetch="one")
    if not sess:
        return jsonify({"error": "Siswa belum memulai ujian ini"}), 404

    questions = query("""
        SELECT q.id, q.content, q.order_num,
               a.option_id as student_option_id,
               ao.label as student_label, ao.content as student_answer,
               ao.is_correct as is_correct,
               co.label as correct_label, co.content as correct_answer
        FROM questions q
        LEFT JOIN answers a ON a.question_id = q.id AND a.session_id = %s
        LEFT JOIN options ao ON ao.id = a.option_id
        LEFT JOIN options co ON co.question_id = q.id AND co.is_correct = true
        WHERE q.exam_id = %s
        ORDER BY q.order_num, q.created_at
    """, (sess["id"], exam_id))

    return jsonify({"questions": [dict(q) for q in questions]})

@app.route("/api/sessions/<session_id>/detail", methods=["GET"])
@require_guru
def session_detail(session_id):
    """Detail jawaban siswa per soal untuk guru."""
    sess = query("SELECT * FROM exam_sessions WHERE id=%s", (session_id,), fetch="one")
    if not sess:
        return jsonify({"error": "Sesi tidak ditemukan"}), 404

    questions = query("""
        SELECT q.id, q.content, q.order_num,
               a.option_id as student_option_id,
               ao.label as student_label, ao.content as student_answer,
               ao.is_correct as is_correct,
               co.label as correct_label, co.content as correct_answer
        FROM questions q
        LEFT JOIN answers a ON a.question_id = q.id AND a.session_id = %s
        LEFT JOIN options ao ON ao.id = a.option_id
        LEFT JOIN options co ON co.question_id = q.id AND co.is_correct = true
        WHERE q.exam_id = %s
        ORDER BY q.order_num, q.created_at
    """, (session_id, sess["exam_id"]))

    result = query("SELECT * FROM results WHERE session_id=%s", (session_id,), fetch="one")

    return jsonify({
        "questions": [dict(q) for q in questions],
        "result": dict(result) if result else None,
        "session": dict(sess),
    })

@app.route("/api/exams/<exam_id>/export", methods=["GET"])
@require_guru
def export_results(exam_id):
    """Export nilai ke Excel."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    exam = query("SELECT * FROM exams WHERE id=%s", (exam_id,), fetch="one")
    if not exam:
        return jsonify({"error": "Ujian tidak ditemukan"}), 404

    rows = query("""
        SELECT r.*, u.name as student_name, u.nisn, c.name as class_name,
               es.tab_violations, es.submitted_at, es.auto_submitted
        FROM results r
        JOIN exam_sessions es ON es.id = r.session_id
        JOIN users   u ON r.student_id = u.id
        LEFT JOIN classes c ON u.class_id = c.id
        WHERE r.exam_id=%s
        ORDER BY c.grade, LENGTH(c.id), c.id, u.name
    """, (exam_id,))

    wb  = Workbook()
    ws  = wb.active
    ws.title = "Nilai Ujian"

    green  = PatternFill("solid", fgColor="0F4C35")
    green2 = PatternFill("solid", fgColor="C8E6C9")
    yellow = PatternFill("solid", fgColor="FFF9C4")
    red    = PatternFill("solid", fgColor="FFCDD2")
    gray   = PatternFill("solid", fgColor="F5F4EF")
    thin   = Border(
        left=Side(style='thin',color='CCCCCC'), right=Side(style='thin',color='CCCCCC'),
        top=Side(style='thin',color='CCCCCC'),  bottom=Side(style='thin',color='CCCCCC'))
    center = Alignment(horizontal="center", vertical="center")

    # Header
    ws.merge_cells("A1:I1")
    ws["A1"] = f"HASIL UJIAN — {exam['title']}"
    ws["A1"].font = Font(bold=True, size=13, color="0F4C35")
    ws["A1"].alignment = center
    ws.merge_cells("A2:I2")
    ws["A2"] = f"SMAN 1 Batangan  |  Dicetak: {now().strftime('%d %B %Y %H:%M')}"
    ws["A2"].alignment = center
    ws["A2"].font = Font(italic=True, size=11, color="888888")
    ws.row_dimensions[3].height = 8

    headers = ["No","Kelas","NISN","Nama Siswa","Benar","Salah",
               "Kosong","Total","Nilai","Pelanggaran Tab","Waktu Submit"]
    for col, h in enumerate(headers, 1):
        c = ws.cell(row=4, column=col, value=h)
        c.font = Font(bold=True, color="FFFFFF")
        c.fill = green
        c.alignment = center
        c.border = thin

    for i, r in enumerate(rows, 1):
        score  = float(r["score"])
        fill   = green2 if score>=75 else (yellow if score>=50 else red)
        submit = r["submitted_at"].strftime("%H:%M:%S") if r["submitted_at"] else "-"
        vals   = [i, r["class_name"], r["nisn"] or "-", r["student_name"],
                  r["correct_count"], r["wrong_count"], r["empty_count"],
                  r["total_questions"], f"{score:.1f}",
                  r["tab_violations"], submit]
        for col, val in enumerate(vals, 1):
            c = ws.cell(row=i+4, column=col, value=val)
            c.alignment = Alignment(horizontal="center" if col!=4 else "left",
                                    vertical="center")
            c.border = thin
            if col == 9: c.fill = fill
            elif i % 2 == 0: c.fill = gray

    for col, w in zip("ABCDEFGHIJK",[5,12,14,28,8,8,8,8,10,16,14]):
        ws.column_dimensions[col].width = w
    ws.row_dimensions[4].height = 20

    buf = io.BytesIO()
    wb.save(buf); buf.seek(0)
    fname = f"Nilai_{exam['title'].replace(' ','_')}_{now().strftime('%Y%m%d')}.xlsx"
    return send_file(buf, as_attachment=True, download_name=fname,
                     mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

# ═══════════════════════════════════════════════════════════

# ═══════════════════════════════════════════════════════════
#  CONVERT DOC → JSON (untuk konverter.html)
# ═══════════════════════════════════════════════════════════
@app.route("/api/convert-doc", methods=["POST"])
def convert_doc():
    """Terima file .doc/.docx, kembalikan soal dalam JSON."""
    if "file" not in request.files:
        return jsonify({"error": "File tidak ditemukan"}), 400
    f   = request.files["file"]
    ext = f.filename.rsplit(".", 1)[-1].lower()

    import tempfile, os, re

    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as tmp:
        f.save(tmp.name)
        tmppath = tmp.name

    try:
        if ext == "docx":
            from docx import Document
            doc = Document(tmppath)
            SKIP = ['LEMBAR','PENILAIAN','TAHUN PELAJARAN','SATUAN','MATA PELAJARAN',
                    'KELAS','WAKTU','KURIKULUM','PETUNJUK','BERILAH TANDA','URAIAN','ESSAY']
            questions = []
            for p in doc.paragraphs:
                text = p.text.strip().replace('\t', '    ')
                if not text or len(text) < 5: continue
                if any(s in text.upper().split('\n')[0] for s in SKIP): continue
                if re.search(r'[a-e]\.\s+\w', text):
                    soal_part = re.split(r'\n\s*[a-e]\.| {4,}[a-e]\.', text, 1)[0].strip()
                    pil = {}
                    for m in re.finditer(r'([a-e])\.\s+([^.]+?)(?=\s{3,}[a-e]\.|$|\n)', text):
                        val = m.group(2).strip()
                        if val and len(val) > 1:
                            pil[m.group(1)] = val
                    if len(pil) >= 2:
                        questions.append({'soal': soal_part,
                            'a': pil.get('a',''), 'b': pil.get('b',''),
                            'c': pil.get('c',''), 'd': pil.get('d',''),
                            'e': pil.get('e',''), 'jawaban': ''})

        elif ext == "doc":
            # .doc lama — coba antiword, fallback ke pesan error
            try:
                import subprocess
                result = subprocess.run(['antiword', tmppath],
                    capture_output=True, text=True, timeout=30)
                if result.returncode == 0:
                    questions = _parse_doc_old(tmppath)
                else:
                    return jsonify({
                        "error": "File .doc lama tidak didukung langsung. Silakan buka di Microsoft Word lalu simpan ulang sebagai .docx (File → Save As → Word Document)",
                        "hint": "docx"
                    }), 400
            except FileNotFoundError:
                return jsonify({
                    "error": "File .doc lama tidak bisa diproses. Silakan buka di Word dan simpan sebagai .docx terlebih dahulu.",
                    "hint": "docx"
                }), 400
        else:
            return jsonify({"error": "Format tidak didukung"}), 400

        return jsonify({"questions": questions, "total": len(questions)})

    finally:
        os.unlink(tmppath)


# ═══════════════════════════════════════════════════════════
#  ADMIN — SEMUA UJIAN
# ═══════════════════════════════════════════════════════════
@app.route("/api/admin/exams", methods=["GET"])
@require_admin
def admin_list_exams():
    rows = query("""
        SELECT e.*, u.name as teacher_name, s.name as subject_name,
               COUNT(DISTINCT q.id) as question_count
        FROM exams e
        LEFT JOIN users    u ON u.id=e.teacher_id
        LEFT JOIN subjects s ON s.id=e.subject_id
        LEFT JOIN questions q ON q.exam_id=e.id
        GROUP BY e.id, u.name, s.name
        ORDER BY e.created_at DESC
    """)
    return jsonify([dict(r) for r in rows])

# ═══════════════════════════════════════════════════════════
#  ADMIN — DATA SISWA
# ═══════════════════════════════════════════════════════════

@app.route("/api/admin/siswa", methods=["GET"])
@require_admin
def admin_list_siswa():
    grade    = request.args.get("grade","")
    class_id = request.args.get("class_id","")
    search   = request.args.get("search","")
    page     = int(request.args.get("page",1))
    per_page = int(request.args.get("per_page",20))

    conditions = ["u.role='siswa'"]
    params = []
    if grade:
        conditions.append("c.grade=%s"); params.append(int(grade))
    if class_id:
        conditions.append("u.class_id=%s"); params.append(class_id)
    if search:
        conditions.append("u.name ILIKE %s"); params.append(f"%{search}%")

    where = " AND ".join(conditions)
    total = query(f"""
        SELECT COUNT(*) as n FROM users u
        LEFT JOIN classes c ON u.class_id=c.id
        WHERE {where}
    """, params, fetch="one")["n"]

    rows = query(f"""
        SELECT u.*, c.name as class_name FROM users u
        LEFT JOIN classes c ON u.class_id=c.id
        WHERE {where}
        ORDER BY c.grade, LENGTH(c.id), c.id, u.name
        LIMIT %s OFFSET %s
    """, params+[per_page,(page-1)*per_page])

    return jsonify({
        "data": [dict(r) for r in rows],
        "total": total,
        "page": page,
        "total_pages": max(1,(total+per_page-1)//per_page)
    })

@app.route("/api/admin/siswa/<siswa_id>", methods=["DELETE"])
@require_admin
def admin_delete_siswa(siswa_id):
    # Hapus data terkait dulu sebelum hapus user
    query("DELETE FROM answers WHERE session_id IN (SELECT id FROM exam_sessions WHERE student_id=%s)", (siswa_id,), fetch="none")
    query("DELETE FROM results WHERE session_id IN (SELECT id FROM exam_sessions WHERE student_id=%s)", (siswa_id,), fetch="none")
    query("DELETE FROM exam_sessions WHERE student_id=%s", (siswa_id,), fetch="none")
    query("DELETE FROM users WHERE id=%s AND role='siswa'", (siswa_id,), fetch="none")
    return jsonify({"ok": True})

@app.route("/api/admin/siswa/<siswa_id>", methods=["GET"])
@require_admin
def admin_get_siswa(siswa_id):
    user = query("""SELECT u.*, c.name as class_name 
                    FROM users u LEFT JOIN classes c ON c.id=u.class_id 
                    WHERE u.id=%s""", (siswa_id,), fetch="one")
    if not user:
        return jsonify({"error": "Siswa tidak ditemukan"}), 404
    return jsonify(dict(user))

@app.route("/api/admin/siswa/<siswa_id>", methods=["PATCH"])
@require_admin
def admin_edit_siswa(siswa_id):
    data = request.json or {}
    name     = data.get("name","").strip()
    email    = data.get("email","").strip() or None
    class_id = data.get("class_id") or None
    nisn     = data.get("nisn","").strip() or None
    if not name:
        return jsonify({"error": "Nama wajib diisi"}), 400
    query("""UPDATE users SET name=%s, email=%s, class_id=%s, nisn=%s
             WHERE id=%s AND role='siswa'""",
          (name, email, class_id, nisn, siswa_id), fetch="none")
    user = query("""SELECT u.*, c.name as class_name 
                    FROM users u LEFT JOIN classes c ON c.id=u.class_id 
                    WHERE u.id=%s""", (siswa_id,), fetch="one")
    return jsonify(dict(user))

@app.route("/api/admin/siswa/import", methods=["POST"])
@require_admin
def admin_import_siswa():
    if "file" not in request.files:
        return jsonify({"error": "File tidak ditemukan"}), 400
    f   = request.files["file"]
    ext = f.filename.rsplit(".",1)[-1].lower()
    rows = []

    import io as _io
    if ext == "csv":
        import csv
        raw = f.read().decode("utf-8-sig", errors="ignore")
        for r in csv.DictReader(_io.StringIO(raw)):
            rows.append({k.strip().lower(): str(v).strip() if v else "" for k,v in r.items()})
    elif ext in ("xlsx","xls"):
        from openpyxl import load_workbook
        wb = load_workbook(_io.BytesIO(f.read()), read_only=True, data_only=True)
        ws = wb.active
        hdrs = [str(v).strip().lower() if v is not None else "" for v in next(ws.iter_rows(values_only=True))]
        for row in ws.iter_rows(min_row=2, values_only=True):
            if not any(v for v in row): continue
            rows.append({hdrs[i]: (str(v).strip() if v is not None else "") for i,v in enumerate(row) if i < len(hdrs)})
    else:
        return jsonify({"error": "Format tidak didukung"}), 400

    def _rk(row, *keys):
        for k in keys:
            v = row.get(k.lower())
            if v and str(v).strip(): return str(v).strip()
        return ""

    def _cls(kelas):
        if not kelas: return None
        k = kelas.lower().strip().replace('-','_').replace(' ','_')
        c = query("SELECT id FROM classes WHERE id=%s", (k,), fetch="one")
        if not c: c = query("SELECT id FROM classes WHERE LOWER(REPLACE(REPLACE(name,'-','_'),' ','_'))=%s",(k,),fetch="one")
        return c["id"] if c else None

    saved, errors = 0, []
    for i, row in enumerate(rows, 1):
        try:
            name  = _rk(row,"nama lengkap *","nama lengkap","nama","name","full name")
            nisn  = _rk(row,"nisn")
            kelas = _rk(row,"kelas","class_id","class")
            email = _rk(row,"email")
            if not name or name.lower() in ("nama lengkap *","nama lengkap","nama","name"): continue
            cid = _cls(kelas)
            ex  = None
            if nisn:  ex = query("SELECT id,class_id FROM users WHERE nisn=%s",(nisn,),fetch="one")
            if not ex and email: ex = query("SELECT id,class_id FROM users WHERE email=%s",(email,),fetch="one")
            if not ex and name:  ex = query("SELECT id,class_id FROM users WHERE LOWER(name)=%s AND role='siswa'",(name.lower(),),fetch="one")
            if ex:
                query("UPDATE users SET name=%s,class_id=%s WHERE id=%s",(name,cid or ex.get("class_id"),ex["id"]),fetch="none")
                if nisn: query("UPDATE users SET nisn=%s WHERE id=%s",(nisn,ex["id"]),fetch="none")
            else:
                dummy_gid = f"import_{uuid.uuid4().hex[:12]}"
                query("INSERT INTO users(id,google_id,name,email,role,class_id,nisn,is_active) VALUES(%s,%s,%s,%s,'siswa',%s,%s,true)",
                      (str(uuid.uuid4()),dummy_gid,name,email or None,cid,nisn or None),fetch="none")
            saved += 1
        except Exception as e:
            errors.append(f"Baris {i}: {str(e)}")

    return jsonify({"saved": saved, "total": len(rows), "errors": errors})


@app.route("/api/admin/siswa/template", methods=["GET"])
@require_admin
def admin_siswa_template():
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill
    import io as _io
    wb = Workbook()
    ws = wb.active
    ws.title = "Template Import Siswa"
    headers = ["nama","nisn","kelas","email"]
    labels  = ["Nama Lengkap *","NISN","Kelas (contoh: x_ipa_1)","Email"]
    green = PatternFill("solid", fgColor="0F4C35")
    for col,(h,l) in enumerate(zip(headers,labels),1):
        c = ws.cell(row=1, column=col, value=l)
        c.font = Font(bold=True, color="FFFFFF")
        c.fill = green
    # Contoh data
    examples = [
        ["Ahmad Fauzi","1234567890","x_ipa_1",""],
        ["Siti Rahayu","0987654321","x_ipa_2",""],
    ]
    for r,ex in enumerate(examples,2):
        for c,val in enumerate(ex,1):
            ws.cell(row=r, column=c, value=val)
    ws.column_dimensions['A'].width = 30
    ws.column_dimensions['B'].width = 18
    ws.column_dimensions['C'].width = 20
    ws.column_dimensions['D'].width = 30
    buf = _io.BytesIO()
    wb.save(buf); buf.seek(0)
    return send_file(buf, mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                     as_attachment=True, download_name="template_import_siswa.xlsx")

# ═══════════════════════════════════════════════════════════
#  ADMIN — KELAS DETAIL
# ═══════════════════════════════════════════════════════════

# ── CRUD Kelas ──────────────────────────────────────────────
@app.route("/api/admin/classes", methods=["POST"])
@require_admin
def admin_create_class():
    data = request.json or {}
    cid   = data.get("id","").strip().lower()
    name  = data.get("name","").strip()
    grade = int(data.get("grade", 10))
    major = data.get("major","Umum").strip()
    if not cid or not name:
        return jsonify({"error": "ID dan nama wajib diisi"}), 400
    existing = query("SELECT id FROM classes WHERE id=%s", (cid,), fetch="one")
    if existing:
        return jsonify({"error": "ID kelas sudah ada"}), 409
    query("INSERT INTO classes (id, name, grade, major) VALUES (%s,%s,%s,%s)",
          (cid, name, grade, major), fetch="none")
    cls = query("SELECT * FROM classes WHERE id=%s", (cid,), fetch="one")
    return jsonify(dict(cls)), 201

@app.route("/api/admin/classes/<class_id>", methods=["PATCH"])
@require_admin
def admin_update_class(class_id):
    data  = request.json or {}
    name  = data.get("name","").strip()
    grade = int(data.get("grade", 10))
    major = data.get("major","Umum").strip()
    if not name:
        return jsonify({"error": "Nama wajib diisi"}), 400
    query("UPDATE classes SET name=%s, grade=%s, major=%s WHERE id=%s",
          (name, grade, major, class_id), fetch="none")
    cls = query("SELECT * FROM classes WHERE id=%s", (class_id,), fetch="one")
    return jsonify(dict(cls))

@app.route("/api/admin/classes/<class_id>", methods=["DELETE"])
@require_admin
def admin_delete_class(class_id):
    # Cek masih ada siswa
    count = query("SELECT COUNT(*) as n FROM users WHERE class_id=%s AND role='siswa'",
                  (class_id,), fetch="one")
    if count and count["n"] > 0:
        return jsonify({"error": f"Masih ada {count['n']} siswa di kelas ini"}), 400
    query("DELETE FROM exam_classes WHERE class_id=%s", (class_id,), fetch="none")
    query("DELETE FROM classes WHERE id=%s", (class_id,), fetch="none")
    return jsonify({"ok": True})

@app.route("/api/admin/classes-detail", methods=["GET"])
@require_admin
def admin_classes_detail():
    rows = query("""
        SELECT c.*, COUNT(u.id) as student_count
        FROM classes c
        LEFT JOIN users u ON u.class_id=c.id AND u.role='siswa'
        GROUP BY c.id
        ORDER BY c.grade, LENGTH(c.id), c.id
    """)
    return jsonify([dict(r) for r in rows])

# ═══════════════════════════════════════════════════════════
#  ADMIN — TAMBAH USER MANUAL
# ═══════════════════════════════════════════════════════════

@app.route("/api/admin/users", methods=["POST"])
@require_admin
def admin_create_user():
    data = request.json or {}
    name     = data.get("name","").strip()
    email    = data.get("email","").strip()
    role     = data.get("role","siswa")
    class_id = data.get("class_id") or None
    nisn     = data.get("nisn") or None
    if not name or not email:
        return jsonify({"error":"Nama dan email wajib diisi"}), 400
    existing = query("SELECT id FROM users WHERE email=%s", (email,), fetch="one")
    if existing:
        return jsonify({"error":"Email sudah terdaftar"}), 409
    uid = str(uuid.uuid4())
    dummy_gid = f"manual_{uuid.uuid4().hex[:12]}"
    query("""INSERT INTO users (id, google_id, email, name, role, class_id, nisn, is_active)
             VALUES (%s,%s,%s,%s,%s,%s,%s,true)""",
          (uid, dummy_gid, email, name, role, class_id, nisn), fetch="none")
    user = query("SELECT * FROM users WHERE id=%s", (uid,), fetch="one")
    return jsonify(dict(user)), 201

# ═══════════════════════════════════════════════════════════
#  EXIT CODE — Guru generate kode untuk izinkan siswa keluar
# ═══════════════════════════════════════════════════════════

@app.route("/api/sessions/<session_id>/allow-exit", methods=["POST"])
@require_guru
def allow_exit(session_id):
    """Guru izinkan siswa tertentu keluar dari lock mode."""
    query("UPDATE exam_sessions SET exit_allowed=true WHERE id=%s", (session_id,), fetch="none")
    return jsonify({"ok": True})

@app.route("/api/sessions/<session_id>/check-exit", methods=["GET"])
@require_auth
def check_exit_allowed(session_id):
    """Siswa cek apakah guru sudah izinkan keluar."""
    sess = query("SELECT exit_allowed FROM exam_sessions WHERE id=%s AND student_id=%s",
                 (session_id, request.user_id), fetch="one")
    if not sess:
        return jsonify({"allowed": False}), 404
    return jsonify({"allowed": sess.get("exit_allowed", False) or False})

# ═══════════════════════════════════════════════════════════
#  GURU — KELAS DIAMPU
# ═══════════════════════════════════════════════════════════

@app.route("/api/guru/taught-classes", methods=["GET"])
@require_guru
def get_taught_classes():
    """Ambil kelas yang diampu guru ini."""
    rows = query("""
        SELECT c.* FROM classes c
        JOIN guru_classes gc ON gc.class_id = c.id
        WHERE gc.teacher_id = %s
        ORDER BY c.grade, LENGTH(c.id), c.id
    """, (request.user_id,))
    return jsonify([dict(r) for r in rows])

@app.route("/api/guru/taught-classes", methods=["POST"])
@require_guru
def add_taught_class():
    """Tambah kelas yang diampu."""
    data = request.json or {}
    class_id = data.get("class_id","").strip()
    if not class_id:
        return jsonify({"error": "class_id wajib diisi"}), 400
    # Cek kelas ada
    cls = query("SELECT id FROM classes WHERE id=%s", (class_id,), fetch="one")
    if not cls:
        return jsonify({"error": "Kelas tidak ditemukan"}), 404
    # Insert (ignore kalau sudah ada)
    query("""
        INSERT INTO guru_classes (teacher_id, class_id)
        VALUES (%s, %s) ON CONFLICT DO NOTHING
    """, (request.user_id, class_id), fetch="none")
    return jsonify({"ok": True})

@app.route("/api/guru/taught-classes/<class_id>", methods=["DELETE"])
@require_guru
def remove_taught_class(class_id):
    """Hapus kelas dari daftar diampu."""
    query("DELETE FROM guru_classes WHERE teacher_id=%s AND class_id=%s",
          (request.user_id, class_id), fetch="none")
    return jsonify({"ok": True})

# ═══════════════════════════════════════════════════════════
#  GURU — DATA SISWA & RIWAYAT
# ═══════════════════════════════════════════════════════════

@app.route("/api/guru/siswa", methods=["GET"])
@require_guru
def guru_get_siswa():
    """Semua siswa beserta rata-rata nilai dari ujian guru ini."""
    teacher_id = request.user_id
    # Ambil semua siswa
    rows = query("""
        SELECT u.id, u.name, u.email, u.nisn, u.class_id,
               c.name as class_name
        FROM users u
        LEFT JOIN classes c ON u.class_id=c.id
        WHERE u.role='siswa'
        ORDER BY c.grade, LENGTH(c.id), c.id, u.name
    """)
    students = [dict(r) for r in rows]

    # Ambil rata-rata nilai per siswa dari ujian guru ini
    scores = query("""
        SELECT es.student_id, AVG(r.score) as avg_score
        FROM exam_sessions es
        JOIN exams e ON e.id=es.exam_id AND e.teacher_id=%s
        LEFT JOIN results r ON r.session_id=es.id
        WHERE r.score IS NOT NULL
        GROUP BY es.student_id
    """, (teacher_id,))
    score_map = {str(s["student_id"]): float(s["avg_score"]) for s in scores if s["avg_score"]}

    for s in students:
        s["avg_score"] = score_map.get(str(s["id"]))

    return jsonify({"students": students})

@app.route("/api/guru/siswa/import", methods=["POST"])
@require_guru
def guru_import_siswa():
    """Guru bisa import siswa — tapi tidak bisa tambah manual."""
    if "file" not in request.files:
        return jsonify({"error": "File tidak ditemukan"}), 400
    f = request.files["file"]
    ext = f.filename.rsplit(".",1)[-1].lower()
    rows = []

    if ext == "csv":
        import csv, io as _io
        raw = f.read().decode("utf-8-sig", errors="ignore")
        reader = csv.DictReader(_io.StringIO(raw))
        # Normalisasi header CSV ke lowercase
        raw_rows = list(reader)
        for r in raw_rows:
            rows.append({k.strip().lower(): (str(v).strip() if v else "") for k,v in r.items()})
    elif ext in ("xlsx","xls"):
        import io as _io
        from openpyxl import load_workbook
        # Baca ke BytesIO dulu agar bisa di-seek
        file_bytes = f.read()
        wb = load_workbook(_io.BytesIO(file_bytes), read_only=True, data_only=True)
        ws = wb.active
        # Header: simpan lowercase
        hdrs = []
        for row in ws.iter_rows(min_row=1, max_row=1, values_only=True):
            hdrs = [str(v).strip().lower() if v is not None else "" for v in row]
        # Data
        for row in ws.iter_rows(min_row=2, values_only=True):
            if not any(v for v in row): continue  # skip baris kosong
            rows.append({hdrs[i]: (str(v).strip() if v is not None else "")
                         for i, v in enumerate(row) if i < len(hdrs)})
    else:
        return jsonify({"error": "Format tidak didukung. Gunakan CSV atau Excel"}), 400

    def resolve_key(row, *keys):
        """Ambil nilai dari row — semua key dicari dalam lowercase."""
        for k in keys:
            v = row.get(k.lower())
            if v and str(v).strip():
                return str(v).strip()
        return ""

    def resolve_class(kelas):
        if not kelas: return None
        k = kelas.lower().strip().replace('-','_').replace(' ','_')
        cls = query("SELECT id FROM classes WHERE id=%s", (k,), fetch="one")
        if not cls:
            cls = query("""
                SELECT id FROM classes
                WHERE LOWER(REPLACE(REPLACE(name,'-','_'),' ','_'))=%s
                   OR LOWER(REPLACE(REPLACE(name,' ','_'),'-','_'))=%s
            """, (k, k), fetch="one")
        return cls["id"] if cls else None

    saved, errors = 0, []
    for i, row in enumerate(rows, 1):
        name = ""
        try:
            name  = resolve_key(row,
                "nama lengkap *","nama lengkap","nama","name","full name")
            nisn  = resolve_key(row, "nisn")
            kelas = resolve_key(row, "kelas","class_id","class")
            email = resolve_key(row, "email")

            if not name or name.lower() in ("nama lengkap *","nama lengkap","nama","name","full name"):
                continue

            class_id = resolve_class(kelas)  # None kalau tidak ketemu — tetap import

            # Cek duplikat by email dulu (paling unik)
            existing = None
            if email:
                existing = query("SELECT id,class_id FROM users WHERE email=%s", (email,), fetch="one")
            if not existing and nisn:
                existing = query("SELECT id,class_id FROM users WHERE nisn=%s", (nisn,), fetch="one")
            if not existing:
                existing = query(
                    "SELECT id,class_id FROM users WHERE LOWER(name)=%s AND role='siswa'",
                    (name.lower(),), fetch="one")

            if existing:
                # Update data yang ada
                query("UPDATE users SET name=%s, class_id=%s, is_active=true WHERE id=%s",
                      (name, class_id or existing.get("class_id"), existing["id"]), fetch="none")
                if email:
                    query("UPDATE users SET email=%s WHERE id=%s", (email, existing["id"]), fetch="none")
                if nisn:
                    query("UPDATE users SET nisn=%s WHERE id=%s", (nisn, existing["id"]), fetch="none")
            else:
                # Insert baru
                dummy_gid = f"import_{uuid.uuid4().hex[:12]}"
                query("""INSERT INTO users (id, google_id, name, email, role, class_id, nisn, is_active)
                         VALUES (%s, %s, %s, %s, 'siswa', %s, %s, true)""",
                      (str(uuid.uuid4()), dummy_gid, name, email or None, class_id, nisn or None), fetch="none")
            saved += 1
        except Exception as e:
            errors.append(f"Baris {i} ({name}): {str(e)}")

    return jsonify({"saved": saved, "total": len(rows), "errors": errors[:10]})

@app.route("/api/template-siswa", methods=["GET"])
@require_guru
def download_template_siswa():
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill
    import io as _io
    wb = Workbook()
    ws = wb.active
    ws.title = "Template Siswa"
    green = PatternFill("solid", fgColor="0F4C35")
    headers = ["nama","nisn","kelas","email"]
    labels  = ["Nama Lengkap *","NISN","Kelas (contoh: x_ipa_1)","Email"]
    for col,(h,l) in enumerate(zip(headers,labels),1):
        c = ws.cell(row=1, column=col, value=l)
        c.font = Font(bold=True, color="FFFFFF")
        c.fill = green
    examples = [
        ["Ahmad Fauzi","1234567890","x_ipa_1",""],
        ["Siti Rahayu","0987654321","xi_ipa_2",""],
    ]
    for r,ex in enumerate(examples,2):
        for c,val in enumerate(ex,1):
            ws.cell(row=r, column=c, value=val)
    ws.column_dimensions['A'].width = 30
    ws.column_dimensions['B'].width = 18
    ws.column_dimensions['C'].width = 20
    ws.column_dimensions['D'].width = 30
    buf = _io.BytesIO()
    wb.save(buf); buf.seek(0)
    return send_file(buf, mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                     as_attachment=True, download_name="template_import_siswa.xlsx")

@app.route("/api/siswa/<siswa_id>/history", methods=["GET"])
@require_guru
def siswa_history(siswa_id):
    rows = query("""
        SELECT e.title as exam_title, es.submitted_at, es.created_at,
               r.score, r.correct_count, r.wrong_count
        FROM exam_sessions es
        JOIN exams e ON e.id=es.exam_id
        LEFT JOIN results r ON r.session_id=es.id
        WHERE es.student_id=%s
        ORDER BY es.created_at DESC
    """, (siswa_id,))
    return jsonify({"history": [dict(r) for r in rows]})

# ═══════════════════════════════════════════════════════════
#  TEMPLATE SOAL & EXPORT NILAI
# ═══════════════════════════════════════════════════════════

@app.route("/api/template-soal", methods=["GET"])
@require_guru
def download_template_soal():
    import io as _io
    csv_content = "soal,A,B,C,D,E,jawaban\n"
    csv_content += "Contoh pertanyaan soal di sini...,Pilihan A,Pilihan B,Pilihan C,Pilihan D,Pilihan E,A\n"
    buf = _io.BytesIO(csv_content.encode("utf-8-sig"))
    return send_file(buf, mimetype="text/csv", as_attachment=True, download_name="template_soal.csv")

@app.route("/api/exams/<exam_id>/export-nilai", methods=["GET"])
@require_guru
def export_nilai(exam_id):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill
    import io as _io

    exam = query("SELECT * FROM exams WHERE id=%s", (exam_id,), fetch="one")
    if not exam:
        return jsonify({"error":"Ujian tidak ditemukan"}), 404

    rows = query("""
        SELECT u.name, u.nisn, c.name as class_name,
               r.score, r.correct_count, r.wrong_count, r.empty_count,
               es.submitted_at, es.tab_violations
        FROM exam_sessions es
        JOIN users u ON u.id=es.student_id
        LEFT JOIN classes c ON c.id=u.class_id
        LEFT JOIN results r ON r.session_id=es.id
        WHERE es.exam_id=%s
        ORDER BY c.grade, LENGTH(c.id), c.id, u.name
    """, (exam_id,))

    wb = Workbook()
    ws = wb.active
    ws.title = "Rekap Nilai"
    green = PatternFill("solid", fgColor="0F4C35")
    headers = ["No","Nama","NISN","Kelas","Benar","Salah","Kosong","Nilai","Submit","Pelanggaran"]
    for col, h in enumerate(headers, 1):
        c = ws.cell(row=1, column=col, value=h)
        c.font = Font(bold=True, color="FFFFFF")
        c.fill = green

    for i, r in enumerate(rows, 2):
        r = dict(r)
        ws.cell(row=i, column=1, value=i-1)
        ws.cell(row=i, column=2, value=r.get("name",""))
        ws.cell(row=i, column=3, value=r.get("nisn",""))
        ws.cell(row=i, column=4, value=r.get("class_name",""))
        ws.cell(row=i, column=5, value=r.get("correct_count",0))
        ws.cell(row=i, column=6, value=r.get("wrong_count",0))
        ws.cell(row=i, column=7, value=r.get("empty_count",0))
        ws.cell(row=i, column=8, value=float(r["score"]) if r.get("score") is not None else "")
        ws.cell(row=i, column=9, value=str(r["submitted_at"]) if r.get("submitted_at") else "Belum")
        ws.cell(row=i, column=10, value=r.get("tab_violations",0))

    ws.column_dimensions['B'].width = 30
    ws.column_dimensions['D'].width = 15
    ws.column_dimensions['H'].width = 10
    ws.column_dimensions['I'].width = 20

    buf = _io.BytesIO()
    wb.save(buf); buf.seek(0)
    filename = f"nilai_{exam['title'].replace(' ','_')}.xlsx"
    return send_file(buf, mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                     as_attachment=True, download_name=filename)


# ═══════════════════════════════════════════════════════════
#  DEVICE TRACKING - Siswa tidak bisa ganti device
# ═══════════════════════════════════════════════════════════

@app.route("/api/admin/siswa/<siswa_id>/device", methods=["GET"])
@require_admin
def get_siswa_device(siswa_id):
    user = query("SELECT id, name, device_id, device_info, last_login FROM users WHERE id=%s", (siswa_id,), fetch="one")
    if not user: return jsonify({"error": "Siswa tidak ditemukan"}), 404
    return jsonify(dict(user))

@app.route("/api/admin/siswa/<siswa_id>/reset-device", methods=["POST"])
@require_admin
def reset_siswa_device(siswa_id):
    query("UPDATE users SET device_id=NULL, device_info=NULL WHERE id=%s AND role='siswa'", (siswa_id,), fetch="none")
    return jsonify({"ok": True, "message": "Device berhasil direset"})

@app.route("/api/student/register-device", methods=["POST"])
@require_auth
def register_device():
    data = request.json or {}
    device_id   = data.get("device_id","").strip()
    device_info = data.get("device_info","").strip()
    if not device_id:
        return jsonify({"error": "device_id wajib"}), 400
    user = query("SELECT device_id FROM users WHERE id=%s", (request.user_id,), fetch="one")
    if not user: return jsonify({"error": "User tidak ditemukan"}), 404
    if user.get("device_id") and user["device_id"] != device_id:
        return jsonify({"allowed": False, "error": "Device tidak dikenali. Hubungi admin untuk reset device."}), 403
    query("UPDATE users SET device_id=%s, device_info=%s WHERE id=%s",
          (device_id, device_info, request.user_id), fetch="none")
    return jsonify({"allowed": True})

# ═══════════════════════════════════════════════════════════
#  ADMIN - GURU SUBJECTS & CLASSES (kelola per guru)
# ═══════════════════════════════════════════════════════════

@app.route("/api/admin/guru/<guru_id>/subjects", methods=["GET"])
@require_admin
def admin_get_guru_subjects(guru_id):
    rows = query("SELECT * FROM subjects WHERE teacher_id=%s ORDER BY name", (guru_id,))
    return jsonify([dict(r) for r in rows])

@app.route("/api/admin/guru/<guru_id>/subjects", methods=["POST"])
@require_admin
def admin_add_guru_subject(guru_id):
    data = request.json or {}
    name = data.get("name","").strip()
    if not name: return jsonify({"error": "Nama mapel wajib"}), 400
    existing = query("SELECT id FROM subjects WHERE LOWER(name)=LOWER(%s) AND teacher_id=%s", (name, guru_id), fetch="one")
    if existing: return jsonify({"error": "Mapel sudah ada untuk guru ini"}), 409
    sub = query("INSERT INTO subjects (id, name, teacher_id) VALUES (%s,%s,%s) RETURNING *",
                (str(uuid.uuid4()), name, guru_id), fetch="one")
    return jsonify(dict(sub)), 201

@app.route("/api/admin/guru/<guru_id>/subjects/<subject_id>", methods=["DELETE"])
@require_admin
def admin_delete_guru_subject(guru_id, subject_id):
    used = query("SELECT id FROM exams WHERE subject_id=%s LIMIT 1", (subject_id,), fetch="one")
    if used: return jsonify({"error": "Mapel masih dipakai di ujian"}), 400
    query("DELETE FROM subjects WHERE id=%s AND teacher_id=%s", (subject_id, guru_id), fetch="none")
    return jsonify({"ok": True})

@app.route("/api/admin/guru/<guru_id>/classes", methods=["GET"])
@require_admin
def admin_get_guru_classes(guru_id):
    rows = query("""
        SELECT c.* FROM classes c
        JOIN guru_classes gc ON gc.class_id=c.id
        WHERE gc.teacher_id=%s
        ORDER BY c.grade, LENGTH(c.id), c.id
    """, (guru_id,))
    return jsonify([dict(r) for r in rows])

@app.route("/api/admin/guru/<guru_id>/classes", methods=["POST"])
@require_admin
def admin_add_guru_class(guru_id):
    data = request.json or {}
    class_id = data.get("class_id","").strip()
    if not class_id: return jsonify({"error": "class_id wajib"}), 400
    query("INSERT INTO guru_classes (teacher_id, class_id) VALUES (%s,%s) ON CONFLICT DO NOTHING",
          (guru_id, class_id), fetch="none")
    return jsonify({"ok": True})

@app.route("/api/admin/guru/<guru_id>/classes/<class_id>", methods=["DELETE"])
@require_admin
def admin_remove_guru_class(guru_id, class_id):
    query("DELETE FROM guru_classes WHERE teacher_id=%s AND class_id=%s", (guru_id, class_id), fetch="none")
    return jsonify({"ok": True})

# ═══════════════════════════════════════════════════════════
#  ADMIN - EXAM SETTINGS (aturan penilaian)
# ═══════════════════════════════════════════════════════════

@app.route("/api/admin/exam-settings", methods=["GET"])
@require_admin
def get_exam_settings():
    settings = query("SELECT * FROM exam_settings LIMIT 1", fetch="one")
    return jsonify(dict(settings) if settings else {
        "passing_grade": 75,
        "allow_remedial": True,
        "max_violations": 5,
        "auto_submit_on_violation": True,
        "show_ranking": True,
    })

@app.route("/api/admin/exam-settings", methods=["PATCH"])
@require_admin
def update_exam_settings():
    data = request.json or {}
    existing = query("SELECT id FROM exam_settings LIMIT 1", fetch="one")
    if existing:
        fields = ", ".join(f"{k}=%s" for k in data.keys())
        query(f"UPDATE exam_settings SET {fields} WHERE id=%s",
              list(data.values()) + [existing["id"]], fetch="none")
    else:
        cols = ", ".join(data.keys())
        vals = ", ".join(["%s"]*len(data))
        query(f"INSERT INTO exam_settings ({cols}) VALUES ({vals})", list(data.values()), fetch="none")
    return jsonify({"ok": True})

if __name__ == "__main__":
    app.run(debug=True, host="0.0.0.0", port=5000)
