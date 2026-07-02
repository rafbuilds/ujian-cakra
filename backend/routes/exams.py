# backend/routes/exams.py
from flask import Blueprint, request, jsonify, send_file
import uuid, io
from db import query, count_correct_wrong, log_activity
from auth import require_guru, require_auth, require_admin
from storage import upload_base64, fetch_image_bytes, is_image_ref

exams_bp = Blueprint('exams', __name__)

# ── CRUD Ujian ─────────────────────────────────────────────────
def _active_semester_id():
    """Semester yang sedang aktif (tahun ajaran aktif + semester aktif) UNTUK SEKOLAH INI.
    None kalau tidak ada."""
    row = query("""
        SELECT s.id FROM semesters s
        JOIN academic_years ay ON ay.id = s.academic_year_id
        WHERE s.is_active=true AND ay.is_active=true AND s.school_id=%s LIMIT 1
    """, (request.school_id,), fetch='one')
    return row['id'] if row else None

@exams_bp.route('/api/exams', methods=['GET'])
@require_guru
def get_exams():
    sem_id = request.args.get('semester_id')
    if not sem_id and request.args.get('active_only') == '1':
        sem_id = _active_semester_id()
    params = [request.user_id, request.school_id]
    extra_where = ""
    if sem_id:
        extra_where = " AND e.semester_id=%s"
        params.append(sem_id)
    rows = query(f"""
        SELECT e.*, s.name as subject_name,
               sem.name as semester_name, ay.name as academic_year_name
               , (SELECT COUNT(*) FROM questions q WHERE q.exam_id=e.id) as question_count,
               (SELECT STRING_AGG(c.name,', ') FROM exam_classes ec
                JOIN classes c ON c.id=ec.class_id WHERE ec.exam_id=e.id) as class_names,
               (SELECT json_agg(json_build_object('id', c.id, 'name', c.name, 'grade', c.grade)
                                ORDER BY c.grade, LENGTH(c.id), c.id)
                FROM exam_classes ec JOIN classes c ON c.id=ec.class_id WHERE ec.exam_id=e.id) as classes
        FROM exams e
        LEFT JOIN subjects s ON s.id=e.subject_id
        LEFT JOIN semesters sem ON sem.id=e.semester_id
        LEFT JOIN academic_years ay ON ay.id=sem.academic_year_id
        WHERE e.teacher_id=%s AND e.school_id=%s{extra_where}
        ORDER BY e.created_at DESC
    """, tuple(params))
    result = []
    for r in rows:
        d = dict(r)
        d['classes'] = d['classes'] or []
        result.append(d)
    return jsonify(result)

@exams_bp.route('/api/exams/all-for-proctor', methods=['GET'])
@require_guru
def get_exams_for_proctor():
    """Semua ujian dari semua guru pada periode (tahun ajaran/semester) yang sedang aktif saja —
    untuk dropdown Pengawas Live universal. Ujian dari semester yang sudah 'Selesai' tidak tampil
    di sini lagi (tetap jadi jejak/record, hanya tidak aktif untuk pengawasan live).
    Hanya metadata, TIDAK termasuk soal/kunci jawaban.

    'classes' berisi daftar kelas peserta per ujian (id/name/grade) — dipakai
    frontend Pengawas Live untuk menyusun folder Jenjang → Kelas, supaya guru
    bisa masuk langsung ke kelas yang dia awasi alih-alih milih dari dropdown
    ujian yang datar."""
    active_id = _active_semester_id()
    rows = query("""
        SELECT e.id, e.title, e.status, e.start_at, e.teacher_id,
               (e.teacher_id = %s) as is_mine,
               s.name as subject_name, u.name as teacher_name,
               (SELECT STRING_AGG(c.name,', ') FROM exam_classes ec
                JOIN classes c ON c.id=ec.class_id WHERE ec.exam_id=e.id) as class_names,
               (SELECT json_agg(json_build_object('id', c.id, 'name', c.name, 'grade', c.grade)
                                ORDER BY c.grade, LENGTH(c.id), c.id)
                FROM exam_classes ec JOIN classes c ON c.id=ec.class_id WHERE ec.exam_id=e.id) as classes
        FROM exams e
        LEFT JOIN subjects s ON s.id=e.subject_id
        LEFT JOIN users u ON u.id=e.teacher_id
        WHERE e.status IN ('published','ongoing','finished')
          AND e.semester_id=%s AND e.school_id=%s
        ORDER BY e.start_at DESC
    """, (request.user_id, active_id, request.school_id))
    result = []
    for r in rows:
        d = dict(r)
        d['classes'] = d['classes'] or []
        result.append(d)
    return jsonify(result)

@exams_bp.route('/api/exams', methods=['POST'])
@require_guru
def create_exam():
    data = request.json or {}
    exam_id = str(uuid.uuid4())
    room_id = data.get('room_id') or None

    # Room ujian (misal "TTS 2026/2027") murni struktur organisasi —
    # tidak mengunci jadwal. Cukup pastikan guru memang sudah join room ini.
    semester_id = None
    if room_id:
        room = query("SELECT semester_id FROM room_teachers rt JOIN rooms r ON r.id=rt.room_id "
                     "WHERE rt.room_id=%s AND rt.teacher_id=%s AND r.school_id=%s",
                     (room_id, request.user_id, request.school_id), fetch='one')
        if not room:
            return jsonify({'error': 'Anda belum join room ujian ini'}), 403
        semester_id = room.get('semester_id')

    # Kalau tidak ada room, atau room-nya belum diset semester, fallback ke
    # semester yang sedang aktif — frozen di waktu pembuatan.
    if not semester_id:
        active_sem = query("SELECT id FROM semesters WHERE is_active=true AND school_id=%s LIMIT 1",
                           (request.school_id,), fetch='one')
        semester_id = active_sem['id'] if active_sem else None

    query("""INSERT INTO exams
             (id, teacher_id, subject_id, title, instructions, duration_minutes,
              start_at, status, randomize_questions, randomize_options,
              show_result_after, show_key_after, score_per_correct, room_id, grade, semester_id, school_id)
             VALUES (%s,%s,%s,%s,%s,%s,%s,'draft',%s,%s,%s,%s,%s,%s,%s,%s,%s)""",
          (exam_id, request.user_id,
           data.get('subject_id') or None,
           data.get('title','Ujian Baru'),
           data.get('instructions',''),
           int(data.get('duration_minutes',90)),
           data.get('start_at'),
           data.get('randomize_questions', True),
           data.get('randomize_options', True),
           data.get('show_result_after', True),
           data.get('show_key_after', False),
           data.get('score_per_correct') or None,
           room_id, data.get('grade') or None, semester_id, request.school_id), fetch='none')
    # Assign classes
    for cls in (data.get('class_ids') or []):
        query("INSERT INTO exam_classes (exam_id, class_id) VALUES (%s,%s) ON CONFLICT DO NOTHING",
              (exam_id, cls), fetch='none')
    log_activity(request.user_id, 'UJIAN_BUAT', f"Buat ujian \"{data.get('title','Ujian Baru')}\"",
                request.remote_addr, request.school_id)
    return jsonify({'id': exam_id}), 201

@exams_bp.route('/api/exams/<exam_id>', methods=['GET'])
@require_guru
def get_exam(exam_id):
    """Hanya guru/admin yang bisa melihat soal + kunci jawaban."""
    exam = query("SELECT * FROM exams WHERE id=%s AND school_id=%s", (exam_id, request.school_id), fetch='one')
    if not exam: return jsonify({'error': 'Tidak ditemukan'}), 404
    # Guru hanya bisa akses ujian miliknya; admin (sekolah yang sama) bisa semua
    if request.user_role not in ('admin', 'super_admin') and str(exam['teacher_id']) != request.user_id:
        return jsonify({'error': 'Akses ditolak'}), 403
    questions = query("SELECT * FROM questions WHERE exam_id=%s ORDER BY order_num", (exam_id,))
    all_opts = query("""SELECT o.* FROM options o JOIN questions q ON q.id=o.question_id
                        WHERE q.exam_id=%s ORDER BY o.label""", (exam_id,))
    opts_by_question = {}
    for o in all_opts:
        opts_by_question.setdefault(str(o['question_id']), []).append(dict(o))
    q_list = [{**dict(q), 'options': opts_by_question.get(str(q['id']), [])} for q in questions]
    classes = query("""SELECT c.* FROM classes c JOIN exam_classes ec ON ec.class_id=c.id
                       WHERE ec.exam_id=%s ORDER BY c.name""", (exam_id,))
    return jsonify({**dict(exam), 'questions': q_list, 'classes': [dict(c) for c in classes]})

@exams_bp.route('/api/exams/<exam_id>', methods=['PATCH'])
@require_guru
def update_exam(exam_id):
    # Pastikan exam ini milik sekolah yang sama, dan guru hanya bisa edit
    # ujian miliknya sendiri (admin sekolah yang sama bebas).
    owner_check = "AND teacher_id=%s" if request.user_role not in ('admin', 'super_admin') else ""
    params = [exam_id, request.school_id] + ([request.user_id] if owner_check else [])
    owner = query(f"SELECT id FROM exams WHERE id=%s AND school_id=%s {owner_check}", tuple(params), fetch='one')
    if not owner:
        return jsonify({'error': 'Ujian tidak ditemukan atau bukan milik Anda'}), 403
    data = request.json or {}
    if 'room_id' in data:
        room_id = data['room_id'] or None
        semester_id = None
        if room_id:
            room = query("SELECT semester_id FROM room_teachers rt JOIN rooms r ON r.id=rt.room_id "
                        "WHERE rt.room_id=%s AND rt.teacher_id=%s AND r.school_id=%s",
                        (room_id, request.user_id, request.school_id), fetch='one')
            if not room:
                return jsonify({'error': 'Anda belum join room ujian ini'}), 403
            semester_id = room.get('semester_id')
        if not semester_id:
            active_sem = query("SELECT id FROM semesters WHERE is_active=true AND school_id=%s LIMIT 1",
                               (request.school_id,), fetch='one')
            semester_id = active_sem['id'] if active_sem else None
        query("UPDATE exams SET room_id=%s, semester_id=%s WHERE id=%s", (room_id, semester_id, exam_id), fetch='none')
    allowed = ['title','instructions','duration_minutes','start_at','status',
               'randomize_questions','randomize_options','show_result_after',
               'show_key_after','subject_id','score_per_correct','grade']
    for f in [k for k in data if k in allowed]:
        query(f"UPDATE exams SET {f}=%s WHERE id=%s", (data[f], exam_id), fetch='none')
    if 'class_ids' in data:
        query("DELETE FROM exam_classes WHERE exam_id=%s", (exam_id,), fetch='none')
        for cls in (data['class_ids'] or []):
            query("INSERT INTO exam_classes (exam_id, class_id) VALUES (%s,%s) ON CONFLICT DO NOTHING",
                  (exam_id, cls), fetch='none')
    if data.get('status') == 'published':
        title = query("SELECT title FROM exams WHERE id=%s", (exam_id,), fetch='one')
        log_activity(request.user_id, 'UJIAN_PUBLISH',
                     f"Publish ujian \"{(title or {}).get('title','')}\"", request.remote_addr, request.school_id)
    return jsonify({'ok': True})

@exams_bp.route('/api/exams/<exam_id>', methods=['DELETE'])
@require_guru
def delete_exam(exam_id):
    query("DELETE FROM exams WHERE id=%s AND teacher_id=%s AND school_id=%s",
          (exam_id, request.user_id, request.school_id), fetch='none')
    return jsonify({'ok': True})

# ── Questions ──────────────────────────────────────────────────
@exams_bp.route('/api/questions/check-similar', methods=['POST'])
@require_guru
def check_similar_questions():
    """Ingatkan guru kalau pernah membuat soal dengan isi mirip (di soal
    miliknya sendiri saja — guru tidak boleh lihat isi soal guru lain)."""
    from auth import feature_blocked_reason
    blocked = feature_blocked_reason(request.school_id, 'bank_soal')
    if blocked:
        return jsonify({'error': blocked}), 403
    content = (request.json or {}).get('content', '').strip()
    if len(content) < 10:
        return jsonify([])
    try:
        rows = query("""
            SELECT q.id as question_id, q.content, q.exam_id, e.title as exam_title,
                   similarity(q.content, %s) as score
            FROM questions q
            JOIN exams e ON e.id = q.exam_id
            WHERE e.teacher_id = %s AND e.school_id = %s AND similarity(q.content, %s) > 0.35
            ORDER BY score DESC LIMIT 5
        """, (content, request.user_id, request.school_id, content))
        return jsonify([dict(r) for r in rows])
    except Exception:
        return jsonify([])

def _exam_owner_or_404(exam_id):
    """Verifikasi exam_id ada & milik sekolah ini, dan (kalau bukan admin)
    milik guru yang request. Return exam dict atau None."""
    owner_check = "AND teacher_id=%s" if request.user_role not in ('admin', 'super_admin') else ""
    params = [exam_id, request.school_id] + ([request.user_id] if owner_check else [])
    return query(f"SELECT * FROM exams WHERE id=%s AND school_id=%s {owner_check}", tuple(params), fetch='one')

@exams_bp.route('/api/exams/<exam_id>/questions', methods=['POST'])
@require_guru
def add_question(exam_id):
    if not _exam_owner_or_404(exam_id):
        return jsonify({'error': 'Ujian tidak ditemukan atau bukan milik Anda'}), 404
    data = request.json or {}
    q_id  = str(uuid.uuid4())
    total = query("SELECT COUNT(*) as n FROM questions WHERE exam_id=%s", (exam_id,), fetch='one')['n']
    q_type         = data.get('type', 'multiple_choice')
    image_url      = upload_base64(data.get('image_url'), folder='soal')
    attachment_url = upload_base64(data.get('attachment_url'), folder='lampiran')
    audio_url      = upload_base64(data.get('audio_url'), folder='audio')
    max_choices    = data.get('max_choices') if q_type == 'multiple_answer' else None

    query("""INSERT INTO questions
               (id, exam_id, content, image_url, type, attachment_url, audio_url, order_num, score, max_choices)
             VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)""",
          (q_id, exam_id, data.get('content',''), image_url,
           q_type, attachment_url, audio_url,
           total+1, data.get('score',1), max_choices), fetch='none')

    for opt in (data.get('options') or []):
        opt_image = upload_base64(opt.get('image_url'), folder='opsi')
        query("INSERT INTO options (id, question_id, label, content, image_url, is_correct) VALUES (%s,%s,%s,%s,%s,%s)",
              (str(uuid.uuid4()), q_id, opt['label'], opt.get('content',''), opt_image, opt.get('is_correct',False)), fetch='none')

    q    = query("SELECT * FROM questions WHERE id=%s", (q_id,), fetch='one')
    opts = query("SELECT * FROM options WHERE question_id=%s ORDER BY label", (q_id,))
    exam_title = query("SELECT title FROM exams WHERE id=%s", (exam_id,), fetch='one')
    log_activity(request.user_id, 'SOAL_BUAT',
                 f"Tambah soal ke \"{(exam_title or {}).get('title','ujian')}\"", request.remote_addr)
    return jsonify({**dict(q), 'options': [dict(o) for o in opts]}), 201

def _question_owner_or_404(question_id):
    """Verifikasi question_id ada, exam-nya milik sekolah ini, dan (kalau
    bukan admin) milik guru yang request."""
    owner_check = "AND e.teacher_id=%s" if request.user_role not in ('admin', 'super_admin') else ""
    params = [question_id, request.school_id] + ([request.user_id] if owner_check else [])
    return query(f"""SELECT q.* FROM questions q JOIN exams e ON e.id=q.exam_id
                     WHERE q.id=%s AND e.school_id=%s {owner_check}""", tuple(params), fetch='one')

@exams_bp.route('/api/questions/<question_id>', methods=['PATCH'])
@require_guru
def update_question(question_id):
    if not _question_owner_or_404(question_id):
        return jsonify({'error': 'Soal tidak ditemukan atau bukan milik Anda'}), 404
    data = request.json or {}
    if 'content' in data:
        query("UPDATE questions SET content=%s WHERE id=%s", (data['content'], question_id), fetch='none')
    if 'type' in data:
        query("UPDATE questions SET type=%s WHERE id=%s", (data['type'], question_id), fetch='none')
    if 'image_url' in data:
        query("UPDATE questions SET image_url=%s WHERE id=%s",
              (upload_base64(data['image_url'], folder='soal'), question_id), fetch='none')
    if 'attachment_url' in data:
        query("UPDATE questions SET attachment_url=%s WHERE id=%s",
              (upload_base64(data['attachment_url'], folder='lampiran'), question_id), fetch='none')
    if 'audio_url' in data:
        query("UPDATE questions SET audio_url=%s WHERE id=%s",
              (upload_base64(data['audio_url'], folder='audio'), question_id), fetch='none')
    if 'max_choices' in data:
        query("UPDATE questions SET max_choices=%s WHERE id=%s", (data['max_choices'], question_id), fetch='none')
    if 'options' in data:
        query("DELETE FROM options WHERE question_id=%s", (question_id,), fetch='none')
        for opt in data['options']:
            opt_image = upload_base64(opt.get('image_url'), folder='opsi')
            query("INSERT INTO options (id, question_id, label, content, image_url, is_correct) VALUES (%s,%s,%s,%s,%s,%s)",
                  (str(uuid.uuid4()), question_id, opt['label'], opt.get('content',''), opt_image, opt.get('is_correct',False)), fetch='none')
    return jsonify({'ok': True})

@exams_bp.route('/api/questions/<question_id>', methods=['DELETE'])
@require_guru
def delete_question(question_id):
    if not _question_owner_or_404(question_id):
        return jsonify({'error': 'Soal tidak ditemukan atau bukan milik Anda'}), 404
    query("DELETE FROM questions WHERE id=%s", (question_id,), fetch='none')
    return jsonify({'ok': True})

# ── Import Soal ────────────────────────────────────────────────
@exams_bp.route('/api/exams/<exam_id>/import', methods=['POST'])
@require_guru
def import_questions(exam_id):
    import traceback, re
    from auth import feature_blocked_reason
    blocked = feature_blocked_reason(request.school_id, 'bank_soal')
    if blocked:
        return jsonify({'error': blocked}), 403
    if not _exam_owner_or_404(exam_id):
        return jsonify({'error': 'Ujian tidak ditemukan atau bukan milik Anda'}), 404
    try:
        file = request.files.get('file')
        if not file: return jsonify({'error': 'File tidak ada'}), 400
        filename = (file.filename or '').lower()
        imported = 0

        def save_question(content, opts_dict, correct_label, idx):
            if not content.strip(): return 0
            q_id = str(uuid.uuid4())
            query("INSERT INTO questions (id, exam_id, content, order_num, score) VALUES (%s,%s,%s,%s,1)",
                  (q_id, exam_id, content.strip(), idx), fetch='none')
            for label, opt_content in opts_dict.items():
                if not opt_content: continue
                query("INSERT INTO options (id, question_id, label, content, is_correct) VALUES (%s,%s,%s,%s,%s)",
                      (str(uuid.uuid4()), q_id, label, opt_content, label == correct_label.upper()), fetch='none')
            return 1

        if filename.endswith('.docx'):
            from docx import Document
            doc = Document(io.BytesIO(file.read()))
            if doc.tables:
                for table in doc.tables:
                    headers = [c.text.strip().lower() for c in table.rows[0].cells]
                    start = 1 if any(h in ['pertanyaan','soal','question'] for h in headers) else 0
                    for row in table.rows[start:]:
                        cells = [c.text.strip() for c in row.cells]
                        if len(cells) < 3 or not cells[0]: continue
                        opts = {l: v for l, v in zip(['A','B','C','D','E'], cells[1:6]) if v}
                        correct = cells[6].strip().upper() if len(cells) > 6 and cells[6].strip() else 'A'
                        imported += save_question(cells[0], opts, correct, imported+1)
            else:
                current_q, current_opts, current_correct = None, {}, 'A'

                def split_inline_option(text):
                    """
                    Pisahkan soal dan opsi pertama jika nempel di baris yang sama.
                    Contoh: "Musik tradisional ... adalah .... a. saron"
                    Return: (soal_bersih, label_opsi, isi_opsi) atau (text, None, None)
                    """
                    # Cari pola " a. " atau " a) " di dalam teks (bukan di awal)
                    m = re.search(r'\s([a-eA-E])[.)]\s+(.+)$', text)
                    if m:
                        soal_part = text[:m.start()].strip()
                        label = m.group(1).upper()
                        isi   = m.group(2).strip()
                        # Hanya split jika bagian soal mengandung angka/teks soal (bukan baris opsi biasa)
                        if re.search(r'\d|\.{2,}|disebut|adalah|merupakan|fungsi|teknik|tujuan', soal_part, re.I):
                            return soal_part, label, isi
                    return text, None, None

                for para in doc.paragraphs:
                    text = para.text.strip()
                    if not text: continue

                    # Cek kunci jawaban
                    m_key = re.match(r'^(?:jawaban|kunci|answer)\s*[:\-]?\s*([A-E])', text, re.I)
                    if m_key:
                        current_correct = m_key.group(1).upper()
                        continue

                    # Cek baris opsi murni (dimulai dengan a. / b. / A) / B) dll)
                    m_opt = re.match(r'^([a-eA-E])[.)]\s*(.+)', text)
                    if m_opt:
                        current_opts[m_opt.group(1).upper()] = m_opt.group(2).strip()
                        continue

                    # Bukan opsi, berarti kemungkinan baris soal baru
                    # Simpan soal sebelumnya dulu
                    if current_q and current_opts:
                        imported += save_question(current_q, current_opts, current_correct, imported+1)

                    # Strip nomor soal di depan: "1. " / "1) " / "1 "
                    m_q = re.match(r'^\d+[.)]\s*(.*)', text)
                    raw_q = m_q.group(1).strip() if m_q else text

                    # Cek apakah opsi pertama nempel di akhir soal
                    soal_clean, first_label, first_isi = split_inline_option(raw_q)
                    current_q = soal_clean
                    current_opts = {}
                    current_correct = 'A'
                    if first_label and first_isi:
                        current_opts[first_label] = first_isi

                if current_q and current_opts:
                    imported += save_question(current_q, current_opts, current_correct, imported+1)

        elif filename.endswith('.csv'):
            import csv
            content_str = file.read().decode('utf-8-sig', errors='replace')
            for i, row in enumerate(csv.reader(io.StringIO(content_str))):
                if i == 0 or not row or not row[0].strip(): continue
                opts = {l: v.strip() for l, v in zip(['A','B','C','D','E'], row[1:6]) if v.strip()}
                correct = row[6].strip().upper() if len(row) > 6 and row[6].strip() else 'A'
                imported += save_question(row[0], opts, correct, imported+1)

        else:
            from openpyxl import load_workbook
            wb = load_workbook(io.BytesIO(file.read()))
            ws = wb.active
            for i, row in enumerate(ws.iter_rows(min_row=2, values_only=True)):
                if not row or not row[0]: continue
                opts = {l: str(v).strip() for l, v in zip(['A','B','C','D','E'], row[1:6]) if v and str(v).strip()}
                correct = str(row[6]).strip().upper() if len(row) > 6 and row[6] else 'A'
                imported += save_question(str(row[0]), opts, correct, imported+1)

        return jsonify({'ok': True, 'saved': imported, 'total': imported})
    except Exception as e:
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

# ── Pengawas Universal — SATU kode global (diset admin) ────────
@exams_bp.route('/api/exams/<exam_id>/join-proctor', methods=['POST'])
@require_guru
def join_proctor(exam_id):
    code = (request.json or {}).get('code', '').strip()
    if not code:
        return jsonify({'error': 'Kode wajib diisi'}), 400
    exam = query("SELECT id, title, teacher_id FROM exams WHERE id=%s AND school_id=%s",
                (exam_id, request.school_id), fetch='one')
    if not exam:
        return jsonify({'error': 'Ujian tidak ditemukan'}), 404
    if str(exam['teacher_id']) == request.user_id:
        return jsonify({'error': 'Anda pembuat soal ujian ini — sudah otomatis bisa mengawasi tanpa kode'}), 400
    setting = query("SELECT proctor_code FROM exam_settings WHERE school_id=%s LIMIT 1",
                    (request.school_id,), fetch='one')
    master_code = (setting or {}).get('proctor_code')
    if not master_code:
        return jsonify({'error': 'Kode pengawas belum diset admin'}), 400
    if code.strip().upper() != master_code.strip().upper():
        return jsonify({'error': 'Kode pengawas salah'}), 403
    query("""INSERT INTO exam_proctors (exam_id, teacher_id) VALUES (%s,%s)
             ON CONFLICT DO NOTHING""", (exam_id, request.user_id), fetch='none')
    return jsonify({'ok': True, 'exam_id': str(exam['id']), 'title': exam['title']})

@exams_bp.route('/api/exams/proctoring', methods=['GET'])
@require_guru
def list_proctoring_exams():
    """ID ujian yang sudah berhasil dibuka guru ini lewat kode (bukan miliknya sendiri)."""
    rows = query("SELECT exam_id FROM exam_proctors WHERE teacher_id=%s", (request.user_id,))
    return jsonify([str(r['exam_id']) for r in rows])

def _is_exam_proctor(exam, user_id, role):
    """Pemilik soal, admin, atau guru yang sudah unlock kode pengawas — sama
    seperti otorisasi monitor_exam."""
    if role == 'admin' or str(exam['teacher_id']) == user_id:
        return True
    return bool(query("SELECT 1 FROM exam_proctors WHERE exam_id=%s AND teacher_id=%s",
                       (exam['id'], user_id), fetch='one'))

# Catatan soal "Keluar Ujian" darurat (kuota internet habis): kode-nya
# (exam_sessions.exit_code) dibuat OTOMATIS sekali per sesi di start_exam
# (siswa.py) — bukan di-generate manual oleh guru — supaya sudah tersimpan
# di browser siswa sejak awal ujian dan tetap bisa dipakai walau offline
# total. monitor_exam() di bawah ikut mengirim exit_code ini (lewat es.*)
# supaya guru bisa lihat langsung di tabel Monitor Siswa tanpa aksi apa pun.
# HP mati/rusak total malah TIDAK butuh kode sama sekali — siswa cukup login
# ulang di device lain (mis. PC lab), sesi yang sama otomatis dilanjutkan.

# ── Monitor ────────────────────────────────────────────────────
@exams_bp.route('/api/exams/<exam_id>/monitor', methods=['GET'])
@require_guru
def monitor_exam(exam_id):
    exam = query("SELECT * FROM exams WHERE id=%s AND school_id=%s", (exam_id, request.school_id), fetch='one')
    if not exam: return jsonify({'error': 'Tidak ditemukan'}), 404
    if not _is_exam_proctor(exam, request.user_id, request.user_role):
        return jsonify({'error': 'Anda belum menjadi pengawas ujian ini. Masukkan kode pengawas terlebih dahulu.'}), 403

    # Filter opsional ke SATU kelas — dipakai saat guru masuk lewat folder
    # kelas di Pengawas Live (bukan pilih ujian langsung), supaya cuma
    # roster kelas itu yang ditampilkan walau ujiannya lintas-kelas.
    class_id = request.args.get('class_id')

    # Roster LENGKAP siswa di kelas peserta ujian ini (LEFT JOIN exam_sessions)
    # — sebelumnya cuma siswa yang SUDAH mulai yang muncul, jadi siswa yang
    # belum login sama sekali tidak kelihatan di layar pengawas (seolah tidak
    # ada), padahal guru justru perlu tahu itu.
    class_filter = "AND ec.class_id=%s" if class_id else ""
    params = [exam_id, exam_id] + ([class_id] if class_id else [])
    roster = query(f"""
        SELECT u.id as student_id, u.name, c.id as class_id, c.name as class_name,
               es.id as session_id, es.submitted_at, es.tab_violations,
               es.exit_allowed, es.exit_code
        FROM exam_classes ec
        JOIN classes c ON c.id=ec.class_id
        JOIN users u ON u.class_id=c.id AND u.role='siswa'
        LEFT JOIN exam_sessions es ON es.exam_id=%s AND es.student_id=u.id
        WHERE ec.exam_id=%s {class_filter}
        ORDER BY c.grade, LENGTH(c.id), c.id, u.name
    """, params)

    total_q = query("SELECT COUNT(*) as n FROM questions WHERE exam_id=%s", (exam_id,), fetch='one')['n']
    started = [r for r in roster if r['session_id']]
    answered_map = {}
    if started:
        session_ids = [str(r['session_id']) for r in started]
        # Gabungkan jawaban single-choice, esai, dan multi-jawaban — supaya
        # progress siswa yang menjawab soal esai/multi tidak terhitung 0.
        answered_rows = query("""
            SELECT session_id, COUNT(DISTINCT question_id) as n FROM (
                SELECT session_id, question_id FROM answers WHERE session_id=ANY(%s::uuid[]) AND option_id IS NOT NULL
                UNION ALL
                SELECT session_id, question_id FROM essay_answers WHERE session_id=ANY(%s::uuid[])
                UNION ALL
                SELECT session_id, question_id FROM multi_answers WHERE session_id=ANY(%s::uuid[])
            ) ans
            GROUP BY session_id
        """, (session_ids, session_ids, session_ids))
        answered_map = {str(r['session_id']): r['n'] for r in answered_rows}

    students = []
    for r in roster:
        has_session = bool(r['session_id'])
        answered = answered_map.get(str(r['session_id']), 0) if has_session else 0
        status = 'submitted' if r.get('submitted_at') else ('ongoing' if has_session else 'not_started')
        students.append({
            'session_id': str(r['session_id']) if has_session else None,
            'student_id': str(r['student_id']),
            'name': r['name'],
            'class_id': str(r['class_id']),
            'class_name': r['class_name'],
            'status': status,
            'answered': answered,
            'total': total_q,
            'tab_violations': r.get('tab_violations', 0) or 0,
            'submitted_at': r['submitted_at'].isoformat() if r.get('submitted_at') else None,
            'exit_allowed': bool(r.get('exit_allowed')),
            'exit_code': r.get('exit_code'),
        })
    submitted   = sum(1 for s in students if s['status'] == 'submitted')
    not_started = sum(1 for s in students if s['status'] == 'not_started')
    ongoing     = len(students) - submitted - not_started
    return jsonify({
        'exam': dict(exam), 'students': students,
        'submitted': submitted, 'ongoing': ongoing, 'not_started': not_started,
        'total_registered': len(students),
    })

# ── Results ────────────────────────────────────────────────────
@exams_bp.route('/api/exams/<exam_id>/results', methods=['GET'])
@require_guru
def exam_results(exam_id):
    if not query("SELECT 1 FROM exams WHERE id=%s AND school_id=%s", (exam_id, request.school_id), fetch='one'):
        return jsonify({'error': 'Ujian tidak ditemukan'}), 404
    # Filter opsional ke satu kelas — dipakai saat guru masuk lewat folder
    # kelas di Nilai & Rekap (bukan pilih ujian langsung), supaya rekap yang
    # ditampilkan cuma untuk kelas itu walau ujiannya lintas-kelas.
    class_id = request.args.get('class_id')
    class_filter = " AND u.class_id=%s" if class_id else ""
    class_params = [class_id] if class_id else []
    rows = query(f"""
        SELECT es.id as session_id, u.id as student_id, u.name, u.nisn,
               c.name as class_name, es.submitted_at, es.tab_violations,
               r.score, r.correct_count, r.wrong_count, r.empty_count
        FROM exam_sessions es
        JOIN users u ON u.id=es.student_id
        LEFT JOIN classes c ON c.id=u.class_id
        LEFT JOIN results r ON r.session_id=es.id
        WHERE es.exam_id=%s{class_filter} ORDER BY c.name, u.name
    """, [exam_id] + class_params)
    all_rows = [dict(r) for r in rows]
    submitted = [r for r in all_rows if r.get('submitted_at')]
    scores = [float(r['score']) for r in submitted if r.get('score') is not None]
    dist = {'a':0,'b':0,'c':0,'d':0}
    for s in scores:
        if s>=90: dist['a']+=1
        elif s>=75: dist['b']+=1
        elif s>=60: dist['c']+=1
        else: dist['d']+=1
    q_stats = []
    session_scope = f"(SELECT id FROM exam_sessions es JOIN users u ON u.id=es.student_id WHERE es.exam_id=%s{class_filter})"
    questions = query(f"""
        SELECT q.id, q.content,
               COUNT(a.id) FILTER (WHERE o.is_correct=true) as correct,
               COUNT(a.id) as total
        FROM questions q
        LEFT JOIN answers a ON a.question_id=q.id AND a.session_id IN {session_scope}
        LEFT JOIN options o ON o.id=a.option_id
        WHERE q.exam_id=%s GROUP BY q.id, q.content ORDER BY q.order_num
    """, [exam_id] + class_params + [exam_id])
    for q in questions:
        total = q['total'] or 1
        pct = round((q['correct'] or 0)/total*100, 1)
        q_stats.append({'correct': q['correct'] or 0, 'total': total, 'pct': pct})
    return jsonify({
        'results': all_rows,
        'summary': {
            'total_students': len(all_rows),
            'submitted': len(submitted),
            'avg_score': round(sum(scores)/len(scores),1) if scores else None,
            'pass_rate': round(len([s for s in scores if s>=75])/len(scores)*100,1) if scores else 0,
        },
        'score_distribution': dist,
        'question_stats': q_stats,
    })

# ── Session Detail ─────────────────────────────────────────────
@exams_bp.route('/api/exams/<exam_id>/student/<student_id>/detail', methods=['GET'])
@require_guru
def student_exam_detail(exam_id, student_id):
    if not query("SELECT 1 FROM exams WHERE id=%s AND school_id=%s", (exam_id, request.school_id), fetch='one'):
        return jsonify({'error': 'Ujian tidak ditemukan'}), 404
    sess = query("SELECT id FROM exam_sessions WHERE exam_id=%s AND student_id=%s LIMIT 1",
                 (exam_id, student_id), fetch='one')
    if not sess: return jsonify({'error': 'Siswa belum memulai ujian ini'}), 404
    session_id = sess['id']

    questions = query("""
        SELECT q.id, q.content, q.order_num,
               COALESCE(q.type, 'multiple_choice') as type,
               q.image_url, q.attachment_url, q.audio_url,
               a.option_id as student_option_id,
               ao.label as student_label, ao.content as student_answer,
               ao.is_correct as is_correct, ao.image_url as student_option_image,
               -- Agregasi semua kunci jawaban menjadi 1 baris per soal
               (SELECT STRING_AGG(label, ', ' ORDER BY label)
                FROM options WHERE question_id=q.id AND is_correct=true) as correct_label,
               (SELECT STRING_AGG(content, ' / ' ORDER BY label)
                FROM options WHERE question_id=q.id AND is_correct=true) as correct_answer,
               (SELECT image_url FROM options
                WHERE question_id=q.id AND is_correct=true LIMIT 1) as correct_option_image
        FROM questions q
        LEFT JOIN answers a ON a.question_id=q.id AND a.session_id=%s
        LEFT JOIN options ao ON ao.id=a.option_id
        WHERE q.exam_id=%s ORDER BY q.order_num, q.created_at
    """, (session_id, exam_id))

    q_list = [dict(q) for q in questions]

    # Ambil essay answers (camera_essay) dan gabungkan
    essay_debug = {'error': None, 'count': 0, 'session_id': str(session_id)}
    try:
        try:
            essays = query("""
                SELECT question_id, essay_text, photo_b64,
                       score as teacher_score, teacher_note
                FROM essay_answers WHERE session_id=%s
            """, (session_id,))
        except Exception:
            try:
                essays = query("""
                    SELECT question_id, essay_text, photo_b64,
                           score as teacher_score
                    FROM essay_answers WHERE session_id=%s
                """, (session_id,))
            except Exception:
                # score / teacher_note belum ada — ambil kolom dasar saja
                essays = query("""
                    SELECT question_id, essay_text, photo_b64
                    FROM essay_answers WHERE session_id=%s
                """, (session_id,))
        essay_debug['count'] = len(essays)
        essay_map = {str(e['question_id']): dict(e) for e in essays}
        for q in q_list:
            qid = str(q['id'])
            # Attach essay jika tipe soal camera_essay ATAU ada data di essay_answers
            if q.get('type') == 'camera_essay' or qid in essay_map:
                essay = essay_map.get(qid, {})
                q['essay_text']    = essay.get('essay_text') or ''
                q['photo_b64']     = essay.get('photo_b64') or ''
                q['teacher_score'] = essay.get('teacher_score')
                q['teacher_note']  = essay.get('teacher_note') or ''
                # Paksa tipe agar frontend render dengan benar
                if qid in essay_map:
                    q['type'] = 'camera_essay'
    except Exception as ex:
        essay_debug['error'] = str(ex)  # tampilkan error, jangan diam

    # Ambil multi_answers dan gabungkan
    try:
        multis = query("""
            SELECT ma.question_id, STRING_AGG(o.label, ', ' ORDER BY o.label) as student_answer,
                   ARRAY_AGG(ma.option_id ORDER BY ma.option_id) as picked
            FROM multi_answers ma
            JOIN options o ON o.id=ma.option_id
            WHERE ma.session_id=%s
            GROUP BY ma.question_id
        """, (session_id,))
        multi_map = {str(m['question_id']): dict(m) for m in multis}
        correct_multi = query("""
            SELECT question_id, STRING_AGG(label, ', ' ORDER BY label) as correct_answer,
                   ARRAY_AGG(id ORDER BY id) as correct_ids
            FROM options WHERE question_id IN (
                SELECT id FROM questions WHERE exam_id=%s AND type='multiple_answer'
            ) AND is_correct=true GROUP BY question_id
        """, (exam_id,))
        correct_multi_map = {str(r['question_id']): dict(r) for r in correct_multi}
        for q in q_list:
            if q['type'] == 'multiple_answer':
                multi = multi_map.get(str(q['id']), {})
                cinfo = correct_multi_map.get(str(q['id']), {})
                q['student_answer'] = multi.get('student_answer', '')
                q['is_correct']     = bool(multi.get('picked')) and multi.get('picked') == cinfo.get('correct_ids')
                q['correct_answer'] = cinfo.get('correct_answer', '—')
    except Exception:
        pass

    return jsonify({'questions': q_list, '_essay_debug': essay_debug})

# ── Debug: cek essay answers langsung ─────────────────────────
@exams_bp.route('/api/debug/essays/<session_id>', methods=['GET'])
@require_guru
def debug_essays(session_id):
    """Endpoint debug — cek isi essay_answers untuk session tertentu."""
    try:
        rows = query("""
            SELECT ea.id, ea.question_id, ea.essay_text,
                   LENGTH(ea.photo_b64) as photo_size,
                   ea.score, ea.submitted_at,
                   q.type as question_type, q.content as question_content
            FROM essay_answers ea
            LEFT JOIN questions q ON q.id = ea.question_id
            WHERE ea.session_id=%s
        """, (session_id,))
        return jsonify({
            'session_id': session_id,
            'count': len(rows),
            'records': [dict(r) for r in rows]
        })
    except Exception as ex:
        return jsonify({'error': str(ex), 'hint': 'Kemungkinan tabel essay_answers belum dibuat'}), 500

# ── Sessions ───────────────────────────────────────────────────
@exams_bp.route('/api/sessions/<session_id>/detail', methods=['GET'])
@require_guru
def session_detail(session_id):
    sess = query("SELECT * FROM exam_sessions WHERE id=%s", (session_id,), fetch='one')
    if not sess: return jsonify({'error': 'Tidak ditemukan'}), 404
    questions = query("""
        SELECT q.id, q.content, q.order_num, a.option_id as student_option_id,
               ao.label as student_label, ao.content as student_answer,
               ao.is_correct, co.label as correct_label, co.content as correct_answer
        FROM questions q
        LEFT JOIN answers a ON a.question_id=q.id AND a.session_id=%s
        LEFT JOIN options ao ON ao.id=a.option_id
        LEFT JOIN options co ON co.question_id=q.id AND co.is_correct=true
        WHERE q.exam_id=%s ORDER BY q.order_num
    """, (session_id, sess['exam_id']))
    result = query("SELECT * FROM results WHERE session_id=%s", (session_id,), fetch='one')
    return jsonify({'questions': [dict(q) for q in questions],
                    'result': dict(result) if result else None, 'session': dict(sess)})

# ── Finish & Publish ───────────────────────────────────────────
@exams_bp.route('/api/exams/<exam_id>/finish', methods=['POST'])
@require_guru
def finish_exam(exam_id):
    exam = query("SELECT * FROM exams WHERE id=%s AND school_id=%s", (exam_id, request.school_id), fetch='one')
    if not exam: return jsonify({'error': 'Ujian tidak ditemukan'}), 404
    if not _is_exam_proctor(exam, request.user_id, request.user_role):
        return jsonify({'error': 'Anda bukan pengawas ujian ini'}), 403

    # Auto-submit semua siswa yang belum mengumpulkan — dulu cuma ubah status
    # ujian jadi 'finished' tanpa benar-benar menutup sesi siswa, padahal
    # dialog konfirmasinya sudah bilang "siswa yang belum submit akan
    # di-auto-submit". Sekarang benar-benar dihitung & disubmit, pakai
    # count_correct_wrong yang sama dengan submit_exam biasa (termasuk
    # soal multiple_answer).
    ongoing = query("SELECT id FROM exam_sessions WHERE exam_id=%s AND submitted_at IS NULL", (exam_id,))
    total_q = query("SELECT COUNT(*) as n FROM questions WHERE exam_id=%s", (exam_id,), fetch='one')['n']
    spc = float(exam.get('score_per_correct') or (100.0/total_q if total_q else 0))
    for sess in ongoing:
        sid = str(sess['id'])
        correct, wrong = count_correct_wrong(sid, exam_id)
        empty = total_q - correct - wrong
        score = round(correct * spc, 2)
        query("UPDATE exam_sessions SET submitted_at=NOW(), auto_submitted=true WHERE id=%s", (sid,), fetch='none')
        existing = query("SELECT id FROM results WHERE session_id=%s", (sid,), fetch='one')
        if existing:
            query("UPDATE results SET score=%s,correct_count=%s,wrong_count=%s,empty_count=%s WHERE session_id=%s",
                  (score, correct, wrong, empty, sid), fetch='none')
        else:
            query("INSERT INTO results (id,session_id,score,correct_count,wrong_count,empty_count) VALUES (%s,%s,%s,%s,%s,%s)",
                  (str(uuid.uuid4()), sid, score, correct, wrong, empty), fetch='none')

    query("UPDATE exams SET status='finished' WHERE id=%s", (exam_id,), fetch='none')
    log_activity(request.user_id, 'UJIAN_SELESAI', f"Selesaikan ujian \"{exam.get('title','')}\"",
                request.remote_addr, request.school_id)
    return jsonify({'ok': True, 'submitted': len(ongoing)})

_QTYPE_LABEL = {'multiple_choice': 'Pilihan Ganda', 'multiple_answer': 'Jawaban Ganda',
                'yes_no': 'Ya/Tidak', 'audio': 'Audio', 'camera_essay': 'Kamera + Uraian'}

def _doc_add_image(doc, data_url, width=2.2):
    from docx.shared import Inches
    img_bytes = fetch_image_bytes(data_url)
    if not img_bytes:
        doc.add_paragraph('[Gambar tidak bisa ditampilkan]')
        return
    try:
        doc.add_picture(io.BytesIO(img_bytes), width=Inches(width))
    except Exception:
        doc.add_paragraph('[Gambar tidak bisa ditampilkan]')

def _doc_write_questions(doc, questions):
    """Tulis daftar soal (pertanyaan+pilihan+kunci+gambar) ke dokumen Word
    yang sedang dibangun — dipakai oleh export_soal (1 ujian) dan
    export_soal_bulk (gabungan beberapa ujian)."""
    from docx.shared import Pt
    if not questions:
        doc.add_paragraph('(Belum ada soal)')
        return
    q_ids = [str(q['id']) for q in questions]
    all_opts = query("SELECT * FROM options WHERE question_id=ANY(%s::uuid[]) ORDER BY label", (q_ids,))
    opts_by_q = {}
    for o in all_opts:
        opts_by_q.setdefault(str(o['question_id']), []).append(o)

    for i, q in enumerate(questions, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {q['content']}")
        run.bold = True
        run.font.size = Pt(12)
        doc.add_paragraph(f"Tipe: {_QTYPE_LABEL.get(q.get('type'), 'Pilihan Ganda')}").runs[0].italic = True
        if is_image_ref(q.get('image_url')):
            _doc_add_image(doc, q['image_url'])
        if is_image_ref(q.get('attachment_url')):
            _doc_add_image(doc, q['attachment_url'])
        for o in opts_by_q.get(str(q['id']), []):
            mark = '  ✓ (KUNCI)' if o.get('is_correct') else ''
            doc.add_paragraph(f"   {o['label']}. {o.get('content') or ''}{mark}")
            if is_image_ref(o.get('image_url')):
                _doc_add_image(doc, o['image_url'], width=1.8)
        doc.add_paragraph('')

# ── Download Soal (Word) — 1 ujian, semua atau hanya soal terpilih ──
@exams_bp.route('/api/exams/<exam_id>/export-soal', methods=['GET'])
@require_guru
def export_soal(exam_id):
    from auth import feature_blocked_reason
    blocked = feature_blocked_reason(request.school_id, 'export')
    if blocked:
        return jsonify({'error': blocked}), 403
    exam = query("SELECT * FROM exams WHERE id=%s AND school_id=%s", (exam_id, request.school_id), fetch='one')
    if not exam: return jsonify({'error': 'Ujian tidak ditemukan'}), 404
    if request.user_role not in ('admin', 'super_admin') and str(exam['teacher_id']) != request.user_id:
        return jsonify({'error': 'Akses ditolak'}), 403

    qids_param = request.args.get('question_ids')
    if qids_param:
        id_list = [x for x in qids_param.split(',') if x]
        if not id_list:
            return jsonify({'error': 'question_ids kosong'}), 400
        questions = query("SELECT * FROM questions WHERE exam_id=%s AND id=ANY(%s::uuid[]) ORDER BY order_num",
                          (exam_id, id_list))
    else:
        questions = query("SELECT * FROM questions WHERE exam_id=%s ORDER BY order_num", (exam_id,))
    if not questions:
        return jsonify({'error': 'Tidak ada soal untuk diunduh'}), 404

    from docx import Document
    doc = Document()
    doc.add_heading(exam.get('title') or 'Soal', level=1)
    _doc_write_questions(doc, questions)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    safe_name = (exam.get('title') or 'soal').replace('/', '-').replace('\\', '-')
    return send_file(buf, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                     as_attachment=True, download_name=f"Soal_{safe_name}.docx")

# ── Download Soal (Word) — GABUNGAN beberapa ujian sekaligus ────
@exams_bp.route('/api/exams/export-soal-bulk', methods=['GET'])
@require_guru
def export_soal_bulk():
    from auth import feature_blocked_reason
    blocked = feature_blocked_reason(request.school_id, 'export')
    if blocked:
        return jsonify({'error': blocked}), 403
    ids_param = request.args.get('exam_ids', '')
    id_list = [x for x in ids_param.split(',') if x]
    if not id_list:
        return jsonify({'error': 'exam_ids wajib'}), 400
    exam_rows = query("SELECT * FROM exams WHERE id=ANY(%s::uuid[]) AND school_id=%s ORDER BY title",
                      (id_list, request.school_id))
    if request.user_role not in ('admin', 'super_admin'):
        exam_rows = [e for e in exam_rows if str(e['teacher_id']) == request.user_id]
    if not exam_rows:
        return jsonify({'error': 'Tidak ada ujian yang bisa diakses'}), 404

    from docx import Document
    doc = Document()
    for idx, exam in enumerate(exam_rows):
        if idx > 0:
            doc.add_page_break()
        doc.add_heading(exam.get('title') or 'Soal', level=1)
        questions = query("SELECT * FROM questions WHERE exam_id=%s ORDER BY order_num", (exam['id'],))
        _doc_write_questions(doc, questions)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return send_file(buf, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                     as_attachment=True, download_name='Soal_Gabungan.docx')

@exams_bp.route('/api/exams/<exam_id>/publish-results', methods=['POST'])
@require_guru
def publish_results(exam_id):
    exam = query("SELECT * FROM exams WHERE id=%s AND school_id=%s", (exam_id, request.school_id), fetch='one')
    if not exam: return jsonify({'error': 'Ujian tidak ditemukan'}), 404
    if not _is_exam_proctor(exam, request.user_id, request.user_role):
        return jsonify({'error': 'Anda bukan pengawas ujian ini'}), 403
    query("UPDATE exams SET status='published', show_result_after=true WHERE id=%s",
          (exam_id,), fetch='none')
    return jsonify({'ok': True})

# ── Grade Essay ────────────────────────────────────────────────
@exams_bp.route('/api/exams/<exam_id>/sessions/<session_id>/grade-essay', methods=['PATCH'])
@require_guru
def grade_essay(exam_id, session_id):
    """Guru beri nilai untuk jawaban esai/foto satu siswa."""
    if not query("SELECT 1 FROM exams WHERE id=%s AND school_id=%s", (exam_id, request.school_id), fetch='one'):
        return jsonify({'error': 'Ujian tidak ditemukan'}), 404
    body  = request.json or {}
    qid   = body.get('question_id')
    score = body.get('score')         # angka 0-100
    note  = body.get('note', '')
    if not qid:
        return jsonify({'error': 'question_id wajib'}), 400
    try:
        # Kolom di tabel: score (bukan teacher_score), teacher_note (ditambah via migration)
        query("""INSERT INTO essay_answers (id, session_id, question_id, score, teacher_note)
                 VALUES (uuid_generate_v4(), %s, %s, %s, %s)
                 ON CONFLICT (session_id, question_id)
                 DO UPDATE SET score=%s, teacher_note=%s""",
              (session_id, qid, score, note, score, note), fetch='none')
        _recalc_with_essay(session_id)
        return jsonify({'ok': True})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

def _recalc_with_essay(session_id):
    """Hitung ulang nilai dengan memasukkan skor esai dari guru."""
    try:
        sess = query("SELECT exam_id FROM exam_sessions WHERE id=%s", (session_id,), fetch='one')
        if not sess: return
        exam_id = sess['exam_id']
        total_q = query("SELECT COUNT(*) as n FROM questions WHERE exam_id=%s", (exam_id,), fetch='one')['n']
        if not total_q: return
        exam    = query("SELECT score_per_correct FROM exams WHERE id=%s", (exam_id,), fetch='one')
        spc     = float(exam.get('score_per_correct') or (100.0 / total_q))

        correct, wrong = count_correct_wrong(session_id, exam_id)
        # Tambah skor esai dari guru
        essay_total = query("""SELECT COALESCE(SUM(score), 0) as s
                               FROM essay_answers WHERE session_id=%s
                               AND score IS NOT NULL""", (session_id,), fetch='one')['s']
        score = round(correct * spc + float(essay_total or 0), 2)
        score = min(100.0, score)
        empty = total_q - correct - wrong
        query("""INSERT INTO results (id, session_id, score, correct_count, wrong_count, empty_count)
                 VALUES (gen_random_uuid(), %s, %s, %s, %s, %s)
                 ON CONFLICT (session_id) DO UPDATE
                 SET score=%s, correct_count=%s, wrong_count=%s, empty_count=%s""",
              (session_id, score, correct, wrong, empty,
               score, correct, wrong, empty), fetch='none')
    except Exception:
        pass

# ── Export Excel ───────────────────────────────────────────────
@exams_bp.route('/api/exams/<exam_id>/export-nilai', methods=['GET'])
@require_guru
def export_nilai(exam_id):
    from auth import feature_blocked_reason
    blocked = feature_blocked_reason(request.school_id, 'export')
    if blocked:
        return jsonify({'error': blocked}), 403
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    exam = query("SELECT * FROM exams WHERE id=%s AND school_id=%s", (exam_id, request.school_id), fetch='one')
    if not exam: return jsonify({'error': 'Ujian tidak ditemukan'}), 404
    rows = query("""
        SELECT u.name, u.nisn, c.name as class_name, es.submitted_at,
               r.score, r.correct_count, r.wrong_count, r.empty_count, es.tab_violations
        FROM exam_sessions es JOIN users u ON u.id=es.student_id
        LEFT JOIN classes c ON c.id=u.class_id
        LEFT JOIN results r ON r.session_id=es.id
        WHERE es.exam_id=%s ORDER BY c.name, u.name
    """, (exam_id,))
    wb = Workbook(); ws = wb.active; ws.title = 'Rekap Nilai'
    header_fill = PatternFill("solid", fgColor="0F4C35")
    headers = ['No','Nama','NISN','Kelas','Waktu Submit','Nilai','Benar','Salah','Kosong','Pelanggaran']
    for j, h in enumerate(headers, 1):
        cell = ws.cell(1, j, h)
        cell.font = Font(bold=True, color='FFFFFF')
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center')
    for i, r in enumerate(rows, 1):
        ws.append([i, r['name'], r['nisn'], r['class_name'],
                   str(r['submitted_at'])[:16] if r['submitted_at'] else 'Belum',
                   float(r['score']) if r['score'] else 0,
                   r['correct_count'], r['wrong_count'], r['empty_count'], r['tab_violations']])
    buf = io.BytesIO(); wb.save(buf); buf.seek(0)
    return send_file(buf, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                     as_attachment=True, download_name=f"nilai_{exam['title'] if exam else 'ujian'}.xlsx")

# ── Export Detail — Jawaban Lengkap per Siswa ──────────────────
@exams_bp.route('/api/exams/<exam_id>/export-detail', methods=['GET'])
@require_guru
def export_detail(exam_id):
    import io
    from auth import feature_blocked_reason
    blocked = feature_blocked_reason(request.school_id, 'export')
    if blocked:
        return jsonify({'error': blocked}), 403
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    from openpyxl.drawing.image import Image as XLImage

    # ── Warna tema ─────────────────────────────────────────
    C_GREEN_DARK  = '0F4C35'
    C_GREEN_MED   = '1D9E75'
    C_GREEN_LIGHT = 'D6F0E6'
    C_RED_LIGHT   = 'FFDADA'
    C_YELLOW      = 'FFF3CD'
    C_BLUE_LIGHT  = 'DBEAFE'
    C_GRAY_HEAD   = 'F0EDE6'
    C_WHITE       = 'FFFFFF'

    thin  = Side(style='thin',  color='CCCCCC')
    thick = Side(style='medium', color='888888')
    def border(t=thin, l=thin, r=thin, b=thin):
        return Border(top=t, left=l, right=r, bottom=b)
    def fill(c): return PatternFill('solid', fgColor=c)
    def font(bold=False, size=11, color='000000', italic=False):
        return Font(bold=bold, size=size, color=color, italic=italic,
                    name='Calibri')
    def align(h='left', v='center', wrap=False):
        return Alignment(horizontal=h, vertical=v, wrap_text=wrap)

    # ── Ambil data ─────────────────────────────────────────
    exam = query("SELECT e.*, u.name as teacher_name, s.name as subject_name "
                 "FROM exams e LEFT JOIN users u ON u.id=e.teacher_id "
                 "LEFT JOIN subjects s ON s.id=e.subject_id "
                 "WHERE e.id=%s AND e.school_id=%s", (exam_id, request.school_id), fetch='one')
    if not exam:
        return jsonify({'error': 'Ujian tidak ditemukan'}), 404
    if request.user_role not in ('admin', 'super_admin') and str(exam['teacher_id']) != request.user_id:
        return jsonify({'error': 'Akses ditolak'}), 403

    sessions = query("""
        SELECT es.id as session_id, es.submitted_at, es.tab_violations,
               es.auto_submitted,
               u.name, u.nisn, c.name as class_name, c.id as class_id,
               r.score, r.correct_count, r.wrong_count, r.empty_count
        FROM exam_sessions es
        JOIN users u ON u.id = es.student_id
        LEFT JOIN classes c ON c.id = u.class_id
        LEFT JOIN results r ON r.session_id = es.id
        WHERE es.exam_id = %s AND es.submitted_at IS NOT NULL
        ORDER BY c.grade, c.name, u.name
    """, (exam_id,))

    questions = query("""
        SELECT q.id, q.content, q.type, q.order_num, q.image_url, q.attachment_url
        FROM questions q WHERE q.exam_id = %s ORDER BY q.order_num
    """, (exam_id,))

    # Bangun option map: {question_id: [options]}
    opt_map = {}
    for q in questions:
        opts = query("SELECT id, label, content, is_correct FROM options "
                     "WHERE question_id=%s ORDER BY label", (q['id'],))
        opt_map[str(q['id'])] = [dict(o) for o in opts]

    # Bangun essay map: {session_id: {question_id: essay}}
    essay_map_all = {}
    for sess in sessions:
        sid = str(sess['session_id'])
        try:
            essays = query("SELECT question_id, essay_text, photo_b64 "
                           "FROM essay_answers WHERE session_id=%s", (sid,))
            essay_map_all[sid] = {str(e['question_id']): dict(e) for e in essays}
        except Exception:
            essay_map_all[sid] = {}

    wb = Workbook()

    # ══════════════════════════════════════════════════════
    # SHEET 1: Rekap Nilai (semua siswa)
    # ══════════════════════════════════════════════════════
    ws_sum = wb.active
    ws_sum.title = 'Rekap Nilai'

    # Judul
    ws_sum.merge_cells('A1:J1')
    ws_sum['A1'] = f'REKAP NILAI UJIAN — {(exam.get("title") or "").upper()}'
    ws_sum['A1'].font      = font(bold=True, size=13, color=C_WHITE)
    ws_sum['A1'].fill      = fill(C_GREEN_DARK)
    ws_sum['A1'].alignment = align('center')
    ws_sum.row_dimensions[1].height = 22

    ws_sum.merge_cells('A2:J2')
    ws_sum['A2'] = f'Guru: {exam.get("teacher_name","")}  |  Mapel: {exam.get("subject_name","")}  |  Durasi: {exam.get("duration_minutes",90)} menit'
    ws_sum['A2'].font      = font(size=10, color='555555', italic=True)
    ws_sum['A2'].alignment = align('center')
    ws_sum.row_dimensions[2].height = 16

    # Header tabel rekap
    h_cols = ['No','Nama Siswa','NISN','Kelas','Waktu Submit',
              'Nilai','Benar','Salah','Kosong','Pelanggaran']
    for j, h in enumerate(h_cols, 1):
        c = ws_sum.cell(3, j, h)
        c.font      = font(bold=True, size=10, color=C_WHITE)
        c.fill      = fill(C_GREEN_MED)
        c.alignment = align('center')
        c.border    = border()
    ws_sum.row_dimensions[3].height = 18

    passing = 75  # default KKM
    for i, s in enumerate(sessions, 1):
        row = i + 3
        score = float(s['score'] or 0)
        lulus = score >= passing
        vals = [i, s['name'], s['nisn'] or '—', s['class_name'] or '—',
                str(s['submitted_at'])[:16] if s['submitted_at'] else '—',
                score, s['correct_count'] or 0,
                s['wrong_count'] or 0, s['empty_count'] or 0,
                s['tab_violations'] or 0]
        for j, v in enumerate(vals, 1):
            c = ws_sum.cell(row, j, v)
            c.alignment = align('center' if j != 2 else 'left')
            c.border    = border()
            if j == 6:  # kolom nilai
                c.font = font(bold=True, color=C_GREEN_DARK if lulus else 'C00000')
                c.fill = fill(C_GREEN_LIGHT if lulus else C_RED_LIGHT)

    # Lebar kolom rekap
    widths_sum = [5, 28, 14, 10, 18, 8, 7, 7, 7, 12]
    for j, w in enumerate(widths_sum, 1):
        ws_sum.column_dimensions[get_column_letter(j)].width = w

    # ══════════════════════════════════════════════════════
    # SHEET PER SISWA
    # ══════════════════════════════════════════════════════
    COLS = ['No', 'Pertanyaan', 'Tipe', 'Jawaban Siswa', 'Jawaban Benar',
            'Status', 'Foto/Essay', 'Nilai Guru', 'Catatan Guru']

    for sess in sessions:
        sid   = str(sess['session_id'])
        sname = (sess['name'] or 'siswa')[:28]
        ws    = wb.create_sheet(title=sname)
        score = float(sess['score'] or 0)
        lulus = score >= passing

        # ── Info siswa (baris 1-6) ───────────────────────
        ws.merge_cells('A1:I1')
        ws['A1'] = f'LEMBAR JAWABAN — {(exam.get("title") or "").upper()}'
        ws['A1'].font      = font(bold=True, size=13, color=C_WHITE)
        ws['A1'].fill      = fill(C_GREEN_DARK)
        ws['A1'].alignment = align('center')
        ws.row_dimensions[1].height = 24

        info = [
            ('Nama Siswa',   sess['name']  or '—',
             'Nilai',        score),
            ('NISN',         sess['nisn']  or '—',
             'Benar',        sess['correct_count'] or 0),
            ('Kelas',        sess['class_name'] or '—',
             'Salah',        sess['wrong_count']  or 0),
            ('Submit',       str(sess['submitted_at'])[:16] if sess['submitted_at'] else '—',
             'Kosong',       sess['empty_count']  or 0),
            ('Pelanggaran',  sess['tab_violations'] or 0,
             'Lulus/Tidak',  'LULUS ✓' if lulus else 'TIDAK LULUS ✗'),
        ]
        for r_off, (lbl1, val1, lbl2, val2) in enumerate(info, 2):
            ws.merge_cells(f'A{r_off}:B{r_off}')
            ws.merge_cells(f'C{r_off}:D{r_off}')
            ws.merge_cells(f'E{r_off}:F{r_off}')
            ws.merge_cells(f'G{r_off}:I{r_off}')

            cl = ws.cell(r_off, 1, lbl1)
            cv = ws.cell(r_off, 3, val1)
            cl2 = ws.cell(r_off, 5, lbl2)
            cv2 = ws.cell(r_off, 7, val2)

            for c in (cl, cl2):
                c.font = font(bold=True, size=10, color='555555')
                c.fill = fill(C_GRAY_HEAD)
                c.alignment = align('right')
                c.border = border()
            for c in (cv, cv2):
                c.font = font(size=10)
                c.alignment = align('left')
                c.border = border()
            # Warna nilai
            if lbl2 == 'Nilai':
                cv2.font = font(bold=True, size=11,
                                color=C_GREEN_DARK if lulus else 'C00000')
            if lbl2 == 'Lulus/Tidak':
                cv2.font = font(bold=True, size=10,
                                color=C_GREEN_DARK if lulus else 'C00000')
                cv2.fill = fill(C_GREEN_LIGHT if lulus else C_RED_LIGHT)

        # ── Header tabel jawaban (baris 8) ───────────────
        HDR_ROW = 8
        ws.row_dimensions[HDR_ROW].height = 20
        for j, h in enumerate(COLS, 1):
            c = ws.cell(HDR_ROW, j, h)
            c.font      = font(bold=True, size=10, color=C_WHITE)
            c.fill      = fill(C_GREEN_DARK)
            c.alignment = align('center', wrap=True)
            c.border    = border(t=thick, l=thick, r=thick, b=thick)

        # ── Ambil jawaban siswa untuk sesi ini ───────────
        ans_rows = query("""
            SELECT a.question_id, a.option_id, o.label, o.content, o.is_correct
            FROM answers a
            LEFT JOIN options o ON o.id = a.option_id
            WHERE a.session_id = %s
        """, (sid,))
        ans_map = {str(r['question_id']): dict(r) for r in ans_rows}
        essay_map = essay_map_all.get(sid, {})

        # ── Tulis soal + jawaban ────────────────────────
        data_row = HDR_ROW + 1
        for q_num, q in enumerate(questions, 1):
            qid   = str(q['id'])
            qtype = q.get('type') or 'multiple_choice'
            ans   = ans_map.get(qid)
            essay = essay_map.get(qid, {})
            opts  = opt_map.get(qid, [])
            correct_opts = [o for o in opts if o.get('is_correct')]
            correct_text = ' / '.join(f"({o['label']}) {o['content']}" for o in correct_opts)

            # Status & warna
            if qtype == 'camera_essay':
                status, row_fill = '📷 Koreksi Manual', C_BLUE_LIGHT
            elif ans and ans.get('is_correct'):
                status, row_fill = '✓ Benar', C_GREEN_LIGHT
            elif ans and ans.get('option_id'):
                status, row_fill = '✗ Salah', C_RED_LIGHT
            else:
                status, row_fill = '— Kosong', C_YELLOW

            # Teks jawaban siswa
            if qtype == 'camera_essay':
                student_ans = essay.get('essay_text') or '(lihat foto)'
            elif ans and ans.get('label'):
                student_ans = f"({ans['label']}) {ans.get('content','')}"
            else:
                student_ans = '—'

            tipe_label = {
                'multiple_choice': 'Pilihan Ganda',
                'camera_essay':    'Esai Foto',
                'essay':           'Uraian',
                'multiple_answer': 'Pilihan Berganda',
                'yes_no':          'Benar/Salah',
            }.get(qtype, qtype)

            vals = [q_num, q['content'] or '', tipe_label,
                    student_ans, correct_text or '—', status,
                    '', '', '']  # foto, nilai guru, catatan guru (kosong)

            ws.row_dimensions[data_row].height = 60
            for j, v in enumerate(vals, 1):
                c = ws.cell(data_row, j, v)
                c.fill      = fill(row_fill)
                c.border    = border()
                c.alignment = align('center' if j in (1, 3, 6, 8) else 'left',
                                    wrap=True)
                c.font      = font(size=10)
                if j == 6:  # status
                    c.font = font(bold=True, size=10,
                                  color=(C_GREEN_DARK if '✓' in status
                                         else ('C00000' if '✗' in status
                                               else '856404')))
                if j == 2:  # pertanyaan
                    c.font = font(size=10)

            # ── Sisipkan foto jika ada — foto jawaban siswa (camera_essay)
            # atau gambar soal sendiri (tipe lain) — kolom G sama, dipakai
            # bergantian karena tidak pernah ada keduanya di baris yang sama.
            photo_b64 = essay.get('photo_b64', '')
            if qtype == 'camera_essay':
                image_ref = photo_b64
            else:
                image_ref = q.get('image_url') if is_image_ref(q.get('image_url')) \
                    else (q.get('attachment_url') if is_image_ref(q.get('attachment_url')) else None)
            if image_ref:
                try:
                    from PIL import Image as PILImage
                    img_bytes = fetch_image_bytes(image_ref)
                    pil_img   = PILImage.open(io.BytesIO(img_bytes))
                    pil_img.thumbnail((180, 140))
                    img_buf = io.BytesIO()
                    pil_img.save(img_buf, format='PNG')
                    img_buf.seek(0)
                    xl_img = XLImage(img_buf)
                    xl_img.width  = 160
                    xl_img.height = 120
                    ws.add_image(xl_img, f'G{data_row}')
                    ws.row_dimensions[data_row].height = 100
                except Exception:
                    ws.cell(data_row, 7, '⚠ Foto tidak dapat dimuat').font = font(italic=True, color='888888')

            data_row += 1

        # ── Kolom "Nilai Guru" dan "Catatan Guru" — bordered kosong ──
        # (sudah terisi kosong, guru isi manual)
        for r in range(HDR_ROW + 1, data_row):
            for j in (8, 9):
                c = ws.cell(r, j)
                c.border = border(t=thin, l=thick, r=thick, b=thin)
                c.fill   = fill('FFFDE7')  # kuning muda — tandai harus diisi

        # ── Lebar kolom per-siswa ────────────────────────
        widths_stu = [5, 50, 14, 36, 36, 14, 22, 12, 30]
        for j, w in enumerate(widths_stu, 1):
            ws.column_dimensions[get_column_letter(j)].width = w

        # ── Freeze header ────────────────────────────────
        ws.freeze_panes = f'A{HDR_ROW + 1}'

    # ── Simpan & kirim ─────────────────────────────────────
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    title_safe = (exam.get('title') or 'ujian').replace('/', '-').replace(' ', '_')
    return send_file(buf,
                     mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                     as_attachment=True,
                     download_name=f'detail_jawaban_{title_safe}.xlsx')

# ── Template Soal ──────────────────────────────────────────────
@exams_bp.route('/api/template-soal', methods=['GET'])
@require_guru
def template_soal():
    from openpyxl import Workbook
    wb = Workbook(); ws = wb.active; ws.title = 'Template Soal'
    ws.append(['Pertanyaan','Pilihan A','Pilihan B','Pilihan C','Pilihan D','Pilihan E','Kunci (A/B/C/D/E)'])
    ws.append(['Contoh: Ibu kota Indonesia adalah?','Jakarta','Surabaya','Bandung','Medan','','A'])
    buf = io.BytesIO(); wb.save(buf); buf.seek(0)
    return send_file(buf, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                     as_attachment=True, download_name='template_soal.xlsx')